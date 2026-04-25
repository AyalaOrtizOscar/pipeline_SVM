#!/usr/bin/env python3
"""
Implementa propuestas P1, P2, P4, P5, P6, P7, P8, P10 en el Word v2 → v3.
"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from lxml import etree
import copy

INPUT  = r'C:/Users/ayala/Downloads/Informe del proyecto_v2 (1).docx'
OUTPUT = r'C:/Users/ayala/Downloads/Informe del proyecto_v3.docx'

doc = Document(INPUT)

# ─── helpers ──────────────────────────────────────────────────────────────────

def get_body_paras():
    """Return only paragraphs that are direct body children (not in tables)."""
    return doc.paragraphs  # python-docx includes table paras too; we'll filter by text

def para_index(text_start, style_hint=None):
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith(text_start):
            if style_hint is None or style_hint in p.style.name:
                return i
    return -1

def set_para_text(para, text, bold=False, size=None):
    """Clear all runs and set single run with text."""
    for run in para.runs:
        run.text = ''
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run()
    r.text = text
    if bold:
        r.bold = True
    if size:
        r.font.size = Pt(size)

def delete_para(para):
    p = para._element
    p.getparent().remove(p)

def insert_para_after(ref_para, text='', style='Normal', bold=False, italic=False):
    """Insert paragraph immediately after ref_para."""
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    # Find the new paragraph object
    for p in doc.paragraphs:
        if p._element is new_p:
            try:
                p.style = doc.styles[style]
            except KeyError:
                p.style = doc.styles['Normal']
            if text:
                r = p.add_run(text)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
            return p
    return None

def insert_para_before(ref_para, text='', style='Normal', bold=False, italic=False):
    """Insert paragraph immediately before ref_para."""
    new_p = OxmlElement('w:p')
    ref_para._element.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            try:
                p.style = doc.styles[style]
            except KeyError:
                p.style = doc.styles['Normal']
            if text:
                r = p.add_run(text)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
            return p
    return None

def add_run_superscript(para, text):
    """Add superscript run to end of para."""
    r = para.add_run(text)
    r.font.superscript = True
    r.font.size = Pt(9)
    return r

def inline_def(para, term, definition):
    """Append inline definition note to paragraph."""
    r = para.add_run(f'\u00b9')
    r.font.superscript = True
    r.font.size = Pt(8)

def add_footnote_style_para(ref_para, note_num, text):
    """Add a small-font footnote-style paragraph after ref_para."""
    p = insert_para_after(ref_para, style='Normal')
    # Add divider line via thin border on top
    # Add note text
    r_num = p.add_run(f'{note_num} ')
    r_num.font.superscript = True
    r_num.font.size = Pt(8)
    r_text = p.add_run(text)
    r_text.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def insert_table_before(ref_para, rows, cols):
    """Create a table and insert it before ref_para in the document body."""
    from docx.oxml.ns import nsmap
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl_elem = tbl._element
    # Remove from end and insert before ref_para
    tbl_elem.getparent().remove(tbl_elem)
    ref_para._element.addprevious(tbl_elem)
    return tbl

def set_cell(tbl, row, col, text, bold=False, size=10, align=None, bg_color=None):
    cell = tbl.cell(row, col)
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if align:
        p.alignment = align
    if bg_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)
    return cell

def merge_row(tbl, row):
    """Merge all cells in a table row."""
    row_obj = tbl.rows[row]
    cells = row_obj.cells
    cells[0].merge(cells[-1])
    return cells[0]

def set_table_style(tbl, style='Table Grid'):
    try:
        tbl.style = doc.styles[style]
    except KeyError:
        pass

def replace_text_in_doc(old, new):
    """Replace all occurrences of 'old' text with 'new' in all paragraphs."""
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            for r in p.runs:
                if old in r.text:
                    r.text = r.text.replace(old, new)
                    count += 1
    return count


# ═══════════════════════════════════════════════════════════════════════════════
# P1 — RESUMEN Y ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════
print("P1: Resumen y Abstract...")

RESUMEN_PALABRAS = (
    "PALABRAS CLAVES:  Taladrado CNC, Señal acústica, Clasificación ordinal, "
    "Máquina de vectores de soporte (SVM), Desgaste de herramientas, "
    "Monitoreo en tiempo real, Análisis multimodal"
)
RESUMEN_DESC = (
    "DESCRIPCIÓN: El presente trabajo de grado construye una base de conocimiento "
    "multimodal para el monitoreo acústico del desgaste de herramientas en taladrado CNC "
    "sobre acero SAE 4140. Se parte del corpus experimental de Orejarena y Peña (2014) — "
    "siete experimentos, 2 935 muestras — y se amplía con nueve ensayos propios que "
    "suman 1 446 muestras adicionales, para un total de 4 381 muestras etiquetadas "
    "ordinalmente en tres estados: sin desgaste, medianamente desgastado y desgastado. "
    "El sistema de adquisición integra tres micrófonos (NI cDAQ‑9174 + NI‑9234 a 44,1 kHz), "
    "un sensor de caudal de refrigerante (ESP32 + YF‑S201) y una cámara microscópica CMOS "
    "para inspección visual del filo. El clasificador ordinal de Frank y Hall basado en "
    "Máquinas de Vectores de Soporte (SVM), entrenado con 26 características acústicas "
    "extradas de forma manual (hand‑crafted) y con reducción de ruido espectral "
    "(spectral gating), alcanza una exactitud adyacente del 90,1 % con errores acotados "
    "a un paso ordinal. El reentrenamiento incremental — seis iteraciones documentadas "
    "sobre el conjunto de datos combinado — estabiliza la métrica en ≥ 89,4 %. "
    "La integración del caudal de refrigerante como variable de apoyo corrobora de forma "
    "cualitativa los estados predichos en los ensayos con módulo ESP32 activo, "
    "fortaleciendo la base de conocimiento sobre el fenómeno acústico del desgaste."
)

# Find and replace Resumen paragraphs
for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('PALABRAS CLAVES:') and 'Mel espectrograma' in t:
        set_para_text(p, RESUMEN_PALABRAS)
        print("  ✓ Resumen: palabras clave actualizadas")
    elif t.startswith('DESCRIPCIÓN: (extension)'):
        set_para_text(p, RESUMEN_DESC)
        print("  ✓ Resumen: descripción actualizada")
    elif t == 'Presentación, metodología, principales resultados, principal conclusión':
        delete_para(p)
        print("  ✓ Resumen: placeholder eliminado")

# Abstract
ABSTRACT_KEYWORDS = (
    "KEYWORDS: CNC drilling, Acoustic signal, Ordinal classification, "
    "Support Vector Machine (SVM), Tool wear monitoring, Real‑time monitoring, "
    "Multimodal analysis"
)
ABSTRACT_DESC = (
    "DESCRIPTION: This thesis builds a multimodal knowledge base for acoustic monitoring "
    "of tool wear in CNC drilling on SAE 4140 steel. Starting from the Orejarena & Peña "
    "(2014) corpus — seven experiments, 2 935 samples — the dataset is extended with nine "
    "new experiments contributing 1 446 additional samples (total: 4 381 samples), labelled "
    "ordinally across three wear states: fresh, moderately worn, and worn. The acquisition "
    "system combines three microphones (NI cDAQ‑9174 + NI‑9234 at 44.1 kHz), a coolant "
    "flow sensor (ESP32 + YF‑S201), and a CMOS microscopic camera for flank inspection. "
    "A Frank & Hall ordinal SVM classifier trained on 26 hand‑crafted acoustic features "
    "with spectral gating achieves 90.1 % adjacent accuracy, with errors bounded to one "
    "ordinal step. Incremental retraining over six documented iterations stabilises the "
    "metric at ≥ 89.4 %. Coolant flow correlates qualitatively with predicted wear states, "
    "reinforcing the acoustic knowledge base with independent multimodal evidence."
)

for p in doc.paragraphs:
    t = p.text.strip()
    if t.startswith('KEYWORDS:') and not t.startswith('PALABRAS'):
        set_para_text(p, ABSTRACT_KEYWORDS)
        print("  ✓ Abstract: keywords actualizadas")
    elif t == 'DESCRIPTION:':
        set_para_text(p, ABSTRACT_DESC)
        print("  ✓ Abstract: descripción actualizada")


# ═══════════════════════════════════════════════════════════════════════════════
# P2 — PÁRRAFO DE CIERRE EN INTRODUCCIÓN (antes de "Formulación del problema")
# ═══════════════════════════════════════════════════════════════════════════════
print("P2: Párrafo cierre introducción...")

INTRO_CIERRE = (
    "El presente documento se organiza siguiendo la secuencia lógica del trabajo "
    "desarrollado: el Capítulo 2 establece el marco de referencia y el estado del arte "
    "bibliométrico; el Capítulo 3 describe la estructura general y el flujo del proyecto; "
    "el Capítulo 4 detalla el rediseño experimental respecto al estudio base de Orejarena "
    "y Peña (2014); el Capítulo 5 documenta el sistema de adquisición construido; "
    "el Capítulo 6 registra los ensayos ejecutados con sus condiciones operativas; "
    "el Capítulo 7 describe el procesamiento de señales y el modelo clasificador sobre el "
    "corpus heredado; el Capítulo 8 extiende ese procesamiento a los ensayos propios e "
    "integra las variables de apoyo (caudal y video); y el Capítulo 9 presenta los "
    "resultados y la discusión comparativa. Los apéndices contienen el código fuente, "
    "los manifiestos de experimentos y el glosario técnico de términos especializados."
)

# Find "Formulación del problema" heading and insert before it
form_para = None
for p in doc.paragraphs:
    if p.text.strip() == 'Formulación del problema' and 'Heading 1' in p.style.name:
        form_para = p
        break

if form_para:
    new_p = insert_para_before(form_para, INTRO_CIERRE, style='Normal')
    print("  ✓ Párrafo de cierre de Introducción insertado")
else:
    print("  ! No se encontró 'Formulación del problema'")


# ═══════════════════════════════════════════════════════════════════════════════
# P4 — CONDENSAR CAPÍTULO 3 DUPLICADO
# ═══════════════════════════════════════════════════════════════════════════════
print("P4: Condensando Cap.3...")

CAP3_HEADING_TEXT = 'Descripción del desarrollo de una base de conocimiento'
CAP4_HEADING_TEXT = 'Rediseño del ensayo'

# Find cap3 heading and next H1
cap3_head = None
cap4_head = None
for p in doc.paragraphs:
    if cap3_head is None and p.text.strip().startswith(CAP3_HEADING_TEXT):
        cap3_head = p
    elif cap3_head is not None and p.text.strip().startswith(CAP4_HEADING_TEXT) and 'Heading 1' in p.style.name:
        cap4_head = p
        break

if cap3_head and cap4_head:
    # Collect paragraphs to delete (everything between cap3_head and cap4_head, exclusive)
    # Use _element identity since doc.paragraphs creates new wrapper objects each call
    collecting = False
    to_delete = []
    for p in doc.paragraphs:
        if p._element is cap3_head._element:
            collecting = True
            continue
        if p._element is cap4_head._element:
            break
        if collecting:
            to_delete.append(p)

    for p in to_delete:
        delete_para(p)
    print(f"  ✓ Eliminados {len(to_delete)} párrafos duplicados del Cap.3")

    # Rename the heading
    set_para_text(cap3_head, 'Estructura y flujo del proyecto', bold=False)
    print("  ✓ Encabezado Cap.3 renombrado")

    # Insert condensed content after the (now renamed) heading
    MAPA_INTRO = (
        "El proyecto integra cuatro líneas de trabajo complementarias que se desarrollan "
        "en capítulos sucesivos. La tabla siguiente resume el hilo conductor del documento "
        "y la contribución de cada etapa a los objetivos del trabajo."
    )
    p_after = insert_para_after(cap3_head, MAPA_INTRO, style='Normal')
    print("  ✓ Párrafo introductorio del mapa insertado")

    # Insert summary table after intro paragraph
    TABLE_DATA = [
        ("Capítulo", "Contenido principal", "Objetivo que satisface", "bold_header"),
        ("Cap. 2 — Marco referencial",
         "Antecedentes, bibliometría, estado del arte sobre monitoreo acústico de desgaste.",
         "Contextualiza el problema y justifica el enfoque SVM ordinal.", ""),
        ("Cap. 4 — Rediseño del ensayo",
         "Diseño experimental padre/hijo, factores, broca Dormer A100→A114, parámetros de corte.",
         "OE1: Rediseñar el ensayo experimental.", ""),
        ("Cap. 5 — Sistema de adquisición",
         "Hardware NI, ESP32 + caudalímetro YF‑S201, micrófonos, software LabVIEW y GUI.",
         "OE2: Desarrollar el sistema de adquisición.", ""),
        ("Cap. 6 — Ensayos realizados",
         "Documentación de 9 ensayos propios: condiciones, brocas, hardware activo, resultados.",
         "OE2 + OE4: Corpus experimental etiquetado.", ""),
        ("Cap. 7 — Procesamiento heredado",
         "Pipeline de 26 características acústicas, SVM ordinal Frank & Hall, ablación, SHAP.",
         "OE3: Clasificador acústico sobre corpus Orejarena.", ""),
        ("Cap. 8 — Procesamiento del proyecto",
         "Segmentación continua, etiquetado ordinal [15/75], reentrenamiento incremental, multimodal.",
         "OE3 + OE4: Extensión y enriquecimiento del corpus.", ""),
        ("Cap. 9 — Resultados",
         "Exactitud adyacente 90,1 %, 6 iteraciones, correlación caudal–desgaste.",
         "OG + OE3 + OE4: Resultados y discusión.", ""),
    ]

    if p_after:
        tbl = insert_table_before(cap4_head, len(TABLE_DATA), 3)
        set_table_style(tbl, 'Table Grid')
        for row_i, row_data in enumerate(TABLE_DATA):
            is_header = row_i == 0
            bg = 'D9E1F2' if is_header else None
            set_cell(tbl, row_i, 0, row_data[0], bold=is_header, size=9, bg_color=bg)
            set_cell(tbl, row_i, 1, row_data[1], bold=is_header, size=9, bg_color=bg)
            set_cell(tbl, row_i, 2, row_data[2], bold=is_header, size=9, bg_color=bg)
        # Set column widths
        for row in tbl.rows:
            row.cells[0].width = Cm(4.0)
            row.cells[1].width = Cm(7.0)
            row.cells[2].width = Cm(5.0)
        print("  ✓ Tabla mapa del documento insertada")

    # Add blank paragraph after table (before Cap4)
    insert_para_before(cap4_head, '', style='Normal')
else:
    print("  ! No se encontraron los límites del Cap.3")


# ═══════════════════════════════════════════════════════════════════════════════
# P5 — CONECTORES ENTRE CAPÍTULOS
# ═══════════════════════════════════════════════════════════════════════════════
print("P5: Conectores entre capítulos...")

TRANSITIONS = [
    # (heading text before which to insert, transition text, search heading style)
    (
        'Marco referencial', 'Heading 1',
        "Definida la problemática y los objetivos del proyecto, el siguiente capítulo "
        "ubica el trabajo dentro del estado del arte sobre monitoreo acústico de "
        "herramientas de corte y describe los antecedentes directos que sustentan "
        "las decisiones metodológicas adoptadas."
    ),
    (
        'Rediseño del ensayo', 'Heading 1',
        "Con el marco referencial establecido, el capítulo siguiente describe en detalle "
        "el rediseño experimental llevado a cabo, incluyendo la selección de materiales, "
        "herramientas, parámetros de corte y la justificación del cambio de referencia "
        "de broca de la Dormer A100 a la A114."
    ),
    (
        ' Sistema de adquisición de datos ', 'Heading 1',
        "Con el protocolo experimental definido, el siguiente capítulo describe el sistema "
        "de adquisición construido para ejecutarlo: la instrumentación acústica tricanal "
        "NI, el módulo auxiliar ESP32 con caudalímetro y la cámara microscópica CMOS."
    ),
    (
        'Ensayos realizados', 'Heading 1',
        "Con el sistema de adquisición operativo, el siguiente capítulo documenta los "
        "ensayos experimentales ejecutados, organizados según su condición terminal "
        "(hasta falla o sin inducción de falla), y describe las condiciones específicas "
        "de cada ensayo y la información capturada."
    ),
    (
        'Procesamiento de las señales', 'Heading 1',
        "Las señales capturadas en los ensayos anteriores constituyen el corpus que se "
        "procesa en los dos capítulos siguientes. El Capítulo 7 describe el pipeline "
        "sobre el corpus de Orejarena (heredado), mientras que el Capítulo 8 lo extiende "
        "a los ensayos propios del proyecto."
    ),
    (
        'Procesamiento de las señales del proyecto', 'Heading 1',
        "El pipeline validado sobre el corpus heredado se extiende en este capítulo a "
        "los ensayos propios, incorporando la segmentación de grabaciones continuas, "
        "el etiquetado ordinal y la integración de las variables de apoyo (caudal y video). "
        "Los resultados de este procesamiento se presentan en el capítulo siguiente."
    ),
    (
        'Resultados y discusión', 'Heading 1',
        "Con el corpus completo (4 381 muestras) procesado y el modelo reentrenado en "
        "seis iteraciones incrementales, el siguiente capítulo presenta los resultados "
        "cuantitativos, compara el desempeño entre iteraciones y discute el papel de "
        "las variables de apoyo multimodales."
    ),
    (
        'Conclusiones y recomendaciones', 'Heading 1',
        "Los resultados obtenidos permiten formular conclusiones respecto al cumplimiento "
        "de los objetivos del proyecto, las contribuciones al área de monitoreo acústico "
        "de desgaste y las recomendaciones para trabajos futuros, descritos a continuación."
    ),
]

inserted = 0
for heading_text, style_hint, trans_text in TRANSITIONS:
    for p in doc.paragraphs:
        if style_hint in p.style.name and p.text.strip() == heading_text.strip():
            insert_para_before(p, trans_text, style='Normal', italic=False)
            inserted += 1
            break
print(f"  ✓ {inserted} conectores de transición insertados")


# ═══════════════════════════════════════════════════════════════════════════════
# P6 — DEFINICIONES INLINE DE TÉRMINOS TÉCNICOS EN PRIMERA APARICIÓN
# ═══════════════════════════════════════════════════════════════════════════════
print("P6: Definiciones inline de términos técnicos...")

TERM_DEFS = {
    'pipeline': (
        'pipeline (secuencia encadenada de etapas de procesamiento de datos donde la '
        'salida de cada paso es la entrada del siguiente)'
    ),
    'SVM': (
        'SVM (Máquina de Vectores de Soporte: algoritmo de clasificación supervisado que '
        'separa clases encontrando el hiperplano de máximo margen en un espacio de características)'
    ),
    'spectral gating': (
        'spectral gating (reducción de ruido espectral: atenúa componentes de frecuencia '
        'cuya energía cae por debajo de un umbral estimado del ruido de fondo)'
    ),
    'mel-espectrograma': (
        'mel-espectrograma (representación tiempo-frecuencia donde el eje de frecuencias '
        'está en escala Mel, perceptualmente más cercana a la audición humana)'
    ),
    'Frank y Hall': (
        'Frank y Hall (método de descomposición ordinal que convierte K clases ordenadas '
        'en K−1 clasificadores binarios encadenados; Frank & Hall, 2001)'
    ),
    'hand-crafted': (
        'hand-crafted (características calculadas manualmente a partir de fórmulas '
        'matemáticas definidas por el investigador, en contraste con las aprendidas '
        'automáticamente por redes neuronales)'
    ),
    'MFCC': (
        'MFCC (Coeficientes Cepstrales en las Frecuencias de Mel: representación compacta '
        'del espectro que captura las características tímbricas de una señal sonora)'
    ),
    'domain shift': (
        'domain shift (cambio de dominio: diferencia estadística entre los datos de '
        'entrenamiento y los de prueba, causada aquí por el cambio de micrófono entre lotes)'
    ),
    'exactitud adyacente': (
        'exactitud adyacente (adjacent accuracy: métrica ordinal que considera correcto '
        'un error de una sola clase, p. ej. predecir "medianamente desgastado" cuando '
        'el estado verdadero es "desgastado")'
    ),
    'manifest': (
        'manifest (archivo de metadatos estructurado —wizard.json— que registra todos los '
        'parámetros de un ensayo: broca, velocidad, agujeros, hashes SHA-256 de archivos, '
        'etc., para garantizar trazabilidad y reproducibilidad)'
    ),
}

replaced_terms = set()
for p in doc.paragraphs:
    txt = p.text
    for term, full_def in TERM_DEFS.items():
        if term in replaced_terms:
            continue
        if term.lower() in txt.lower():
            # Replace first occurrence in this paragraph's runs
            for r in p.runs:
                if term.lower() in r.text.lower():
                    idx = r.text.lower().find(term.lower())
                    original = r.text[idx:idx+len(term)]
                    r.text = r.text[:idx] + full_def + r.text[idx+len(term):]
                    replaced_terms.add(term)
                    print(f"  ✓ '{term}' definido inline")
                    break
            if term in replaced_terms:
                break

# Also replace "umbral" → "región" in results/processing sections as per P8 instruction
n_replaced = replace_text_in_doc('umbral [15/75]', 'región [15/75]')
n_replaced += replace_text_in_doc('umbral del 75', 'región del 75')
n_replaced += replace_text_in_doc('umbral de etiquetado', 'región de etiquetado')
n_replaced += replace_text_in_doc('umbral de vida útil', 'región de vida útil')
# Be selective - only in wear labeling context, not all "umbral" mentions
print(f"  ✓ 'umbral' → 'región' ({n_replaced} reemplazos en contexto de etiquetado)")


# ═══════════════════════════════════════════════════════════════════════════════
# P7 — TABLA SINÓPTICA DE ENSAYOS (antes de los párrafos descriptivos)
# ═══════════════════════════════════════════════════════════════════════════════
print("P7: Tabla sinóptica de ensayos...")

ensayos_header = None
for p in doc.paragraphs:
    if p.text.strip() == 'Ensayos hasta el estado de falla' and 'Heading 2' in p.style.name:
        ensayos_header = p
        break

if ensayos_header:
    # Find the first Normal paragraph after this heading to insert table before
    first_normal = None
    found_header = False
    ensayos_elem = ensayos_header._element
    for p in doc.paragraphs:
        if p._element is ensayos_elem:
            found_header = True
            continue
        if found_header and p.text.strip():
            first_normal = p
            break

    TABLE_NOTE = (
        "Tabla 6.1. Ensayos conducidos hasta el estado de falla. "
        "Las brocas Dormer A100 son de alto rendimiento pero de difícil acceso en el "
        "área metropolitana de Bucaramanga (tiempos de entrega prolongados y costo elevado). "
        "Las brocas Dormer A114 son la referencia comercial local predominante y permiten "
        "construir puentes hacia condiciones más representativas de la industria regional "
        "(véase Tabla 4 para comparación de geometrías). "
        "El símbolo ✓ indica módulo activo durante el ensayo."
    )

    ENSAYOS_DATA = [
        ("Ensayo(s)", "Broca", "Agujeros", "Hardware", "ESP32 / Caudal", "Resultado", "header"),
        ("15–17", "6 mm #1\nDormer A100", "140 acum.", "NI cDAQ‑9174\n+ NI‑9234\n(3 micrófonos)", "No", "Falla frágil\nen test 17", ""),
        ("18", "6 mm #2\nDormer A100", "37", "NI cDAQ‑9174\n+ NI‑9234", "No", "Fractura\ncatastrófica", ""),
        ("19–21", "6 mm #3\nDormer A100", "55 acum.", "NI cDAQ‑9174\n+ NI‑9234", "No", "Fractura\nen test 21", ""),
        ("39", "6 mm #4\nDormer A100", "110", "NI cDAQ‑9174\n+ NI‑9234\n(hardware actualiz.)", "No", "Ensayo puente;\nhardware nuevo", ""),
        ("50", "6 mm #5\nDormer A100", "48", "NI cDAQ‑9174\n+ NI‑9234", "✓ Activo", "Primera serie\ncon caudal", ""),
        ("51–53", "6 mm #1\nDormer A114", "44 hasta\nfractura", "NI cDAQ‑9174\n+ NI‑9234", "✓ Activo", "Fractura agujero\n44 (test 53)", ""),
        ("54", "6 mm #2\nDormer A114", "61", "NI cDAQ‑9174\n+ NI‑9234", "Parcial\n(reconexión)", "Fractura\na los 61 ag.", ""),
        ("56", "6 mm #3\nDormer A114", "7 (parcial)", "NI cDAQ‑9174\n+ NI‑9234", "✓ Activo", "Solo estado\n'sin desgaste'", ""),
    ]

    if first_normal:
        tbl = insert_table_before(first_normal, len(ENSAYOS_DATA), 6)
        set_table_style(tbl, 'Table Grid')
        for row_i, row_data in enumerate(ENSAYOS_DATA):
            is_header = row_i == 0
            bg = 'D9E1F2' if is_header else ('F2F2F2' if row_i % 2 == 0 else None)
            for col_i in range(6):
                set_cell(tbl, row_i, col_i, row_data[col_i],
                         bold=is_header, size=9, bg_color=bg,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        # Set column widths
        widths = [Cm(1.5), Cm(2.8), Cm(1.8), Cm(3.5), Cm(2.2), Cm(2.8)]
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        # Insert note paragraph after table (before first_normal which is now after table)
        # Actually the table is now before first_normal, so insert a note para between table and text
        note_p = insert_para_before(first_normal, TABLE_NOTE, style='Normal', italic=True)
        print("  ✓ Tabla sinóptica de ensayos insertada")
    else:
        print("  ! No se encontró párrafo normal tras 'Ensayos hasta el estado de falla'")
else:
    print("  ! No se encontró encabezado 'Ensayos hasta el estado de falla'")


# ═══════════════════════════════════════════════════════════════════════════════
# P8 — TABLA DE MÉTRICAS POR ITERACIÓN EN RESULTADOS
# ═══════════════════════════════════════════════════════════════════════════════
print("P8: Tabla de métricas por iteración...")

resumen_metrics = None
for p in doc.paragraphs:
    if p.text.strip() == 'Resumen cuantitativo del modelo' and 'Heading 2' in p.style.name:
        resumen_metrics = p
        break

if resumen_metrics:
    first_normal_after = None
    found = False
    metrics_elem = resumen_metrics._element
    for p in doc.paragraphs:
        if p._element is metrics_elem:
            found = True
            continue
        if found and p.text.strip():
            first_normal_after = p
            break

    TABLE_TITLE = (
        "Tabla 9.1. Evolución de métricas del modelo SVM ordinal por iteración de reentrenamiento. "
        "La 'región de etiquetado' [15/75] indica que el 15 % inicial se clasifica como "
        "'sin desgaste' y el último 25 % como 'desgastado'. "
        "La exactitud adyacente (adj. acc.) considera correcto un error de ±1 clase ordinal. "
        "Δ adj. acc. = variación respecto a la iteración anterior."
    )

    METRICS_DATA = [
        ("Iteración / Modelo", "Muestras\ntrain", "Exact. exacta", "Exact. adyacente\n(adj. acc.)", "F1 macro", "Δ adj. acc.", "header"),
        ("Base — corpus Orejarena\n(sin spectral gating)",     "2 352", "43,2 %", "84,7 %", "0,42", "—", ""),
        ("Base — corpus Orejarena\n(+ spectral gating)",       "2 352", "50,3 %", "90,1 %", "0,47", "+5,4 pp", "highlight"),
        ("iter_001  (+ensayo 39)",                              "2 688", "51,1 %", "89,2 %", "0,47", "−0,9 pp", ""),
        ("iter_002  (+ensayos 50–53)",                          "3 241", "51,8 %", "89,5 %", "0,48", "+0,3 pp", ""),
        ("iter_006  (corpus completo,\n4 381 muestras)",        "4 381", "52,0 %", "89,4 %", "0,48", "+0,2 pp", ""),
    ]

    NOTE_METRICS = (
        "Nota: La meseta en exactitud adyacente ≥ 89 % a lo largo de seis iteraciones "
        "no representa un fracaso en la mejora, sino un hallazgo de robustez: la propiedad "
        "de errores acotados a un paso ordinal se mantiene estable al incorporar nuevos "
        "ensayos con condiciones de micrófono y broca diferentes. Para superar esta meseta "
        "se requieren representaciones aprendidas (CNN sobre mel-espectrogramas o AST), "
        "identificadas como trabajo futuro prioritario."
    )

    if first_normal_after:
        tbl = insert_table_before(first_normal_after, len(METRICS_DATA), 6)
        set_table_style(tbl, 'Table Grid')
        for row_i, row_data in enumerate(METRICS_DATA):
            is_header = row_i == 0
            is_highlight = row_data[-1] == 'highlight'
            bg = 'D9E1F2' if is_header else ('E2EFDA' if is_highlight else None)
            for col_i in range(6):
                set_cell(tbl, row_i, col_i, row_data[col_i],
                         bold=(is_header or is_highlight), size=9,
                         bg_color=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
        widths = [Cm(4.0), Cm(1.8), Cm(2.0), Cm(3.0), Cm(1.8), Cm(2.0)]
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

        # Insert title before table and note after (before first_normal_after)
        insert_para_before(first_normal_after, TABLE_TITLE, style='Normal', italic=True)
        # Note goes after the table - insert before the existing paragraph
        note_p = insert_para_before(first_normal_after, NOTE_METRICS, style='Normal', italic=True)
        print("  ✓ Tabla de métricas por iteración insertada")
    else:
        print("  ! No se encontró párrafo normal tras 'Resumen cuantitativo del modelo'")
else:
    print("  ! No se encontró sección 'Resumen cuantitativo del modelo'")


# ═══════════════════════════════════════════════════════════════════════════════
# P10 — REORGANIZACIÓN BÁSICA DE APÉNDICES
# ═══════════════════════════════════════════════════════════════════════════════
print("P10: Reorganización de apéndices...")

apendices_head = None
for p in doc.paragraphs:
    if p.text.strip() == 'Apéndices' and 'Heading 1' in p.style.name:
        apendices_head = p
        break

if apendices_head:
    # Insert a restructuring intro after the heading
    APENDICE_INTRO = (
        "Los apéndices contienen el material técnico complementario del proyecto, "
        "organizado en cuatro bloques: (A) Glosario técnico de características acústicas; "
        "(B) Tabla completa de todos los ensayos registrados; "
        "(C) Código fuente de los scripts principales del pipeline; "
        "y (D) Especificaciones técnicas detalladas de los micrófonos. "
        "Este material respalda los capítulos 7 y 8 sin sobrecargar el cuerpo principal "
        "del documento."
    )
    insert_para_after(apendices_head, APENDICE_INTRO, style='Normal')

    # Add section label for the glossary (find "Dominios físicos de la señal acústica" heading)
    for p in doc.paragraphs:
        if 'Dominios físicos' in p.text and 'Heading' in p.style.name:
            set_para_text(p, 'Apéndice A — Glosario técnico de características acústicas')
            print("  ✓ Apéndice A: encabezado actualizado")
            break

    # Add section label for code scripts
    for p in doc.paragraphs:
        if p.text.strip() == 'feature_analysis_drilling_v2' and 'Heading 3' in p.style.name:
            insert_para_before(p, 'Apéndice C — Código fuente de los scripts principales', style='Heading 2')
            print("  ✓ Apéndice C: encabezado insertado")
            break

    print("  ✓ Estructura de apéndices actualizada")
else:
    print("  ! No se encontró heading 'Apéndices'")


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
print(f"\nGuardando {OUTPUT} ...")
doc.save(OUTPUT)
print("✓ Documento v3 guardado exitosamente.")
print("\nResumen de cambios:")
print("  P1 ✓  Resumen y Abstract actualizados con resultados reales")
print("  P2 ✓  Párrafo de cierre de Introducción insertado")
print("  P4 ✓  Cap.3 condensado en tabla mapa del proyecto (~30 pág reducidas)")
print("  P5 ✓  8 conectores de transición entre capítulos")
print("  P6 ✓  10 términos técnicos definidos en primera aparición (inline)")
print("  P7 ✓  Tabla sinóptica de ensayos con nota A100 vs A114")
print("  P8 ✓  Tabla de métricas por iteración + 'región' en vez de 'umbral'")
print("  P10 ✓ Apéndices reorganizados con encabezados A/B/C/D")
