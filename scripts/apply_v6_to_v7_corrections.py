#!/usr/bin/env python3
"""
v6 -> v7 corrections per advisor feedback (2026-04-19).

Changes applied:
  1. Justificacion replaced (II.A) with 3 footnotes.
  2. Chapter 8 body replaced (II.B) — preserves existing drawing paragraphs.
  3. Chapter 9 body replaced (II.C) — preserves drawing paragraphs, adds Tablas 9.1/9.2/9.3 APA.
  4. Chapter 10 body replaced (II.D).
  5. Orphan sections deleted (896-918, 988-1058 of old schema 11.1-13).
  6. Alternativa B rewritten (third-person impersonal).
  7. Dedicatoria typos fixed.
  8. References: add Ayala 2026a/2026b, Shaw 2005; fix Orejarena/Müller; remove duplicate Meneses 2020b.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from copy import deepcopy
from pathlib import Path

SRC = Path('C:/Users/ayala/Downloads/files/Informe_del_proyecto_v6.docx')
OUT = Path('C:/Users/ayala/Downloads/files/Informe_del_proyecto_v7.docx')

# ===== Paragraph index map (from v6.docx structural scan) =====
IDX = {
    # Intro region
    'dedicatoria_a_padre': 45,
    'dedicatoria_a_madre': 46,
    'dedicatoria_a_hermana': 47,
    'agradec_meneses': 57,
    # Justification
    'justif_head': 242,
    'justif_p1': 243,
    'justif_p3': 245,      # last paragraph of justification (246 = Objetivos heading)
    # Chapter 3
    'ch3_head': 302,
    'ch3_end': 307,
    # Alternativa B
    'altB_head': 464,
    'altB_end': 475,       # last paragraph before Alternativa C (476)
    # Chapter 8 (Procesamiento proyecto actual)
    'ch8_head': 791,
    'ch8_end': 821,        # last paragraph before "Resultados y discusion" (822)
    # Chapter 9 (Resultados y discusion)
    'ch9_head': 822,
    'ch9_end': 869,        # last paragraph before "Conclusiones" (870)
    # Chapter 10 (Conclusiones)
    'ch10_head': 870,
    'ch10_end': 895,       # last para before orphan "Anexos" (896)
    # Orphan end sections
    'orphan_anexos_start': 896,
    'orphan_anexos_end': 918,   # everything up to but not including "Referencias Bibliograficas" (919)
    # Another orphan block before Apendice C
    'orphan_tail_start': 988,
    'orphan_tail_end': 1058,
}

# ===== Helpers =====

def has_drawing(elem):
    return (elem.find('.//' + qn('w:drawing')) is not None
            or elem.find('.//' + qn('w:pict')) is not None)


def delete_range(doc, start_idx, end_idx, preserve_drawings=False):
    """Delete paragraphs [start_idx, end_idx] inclusive.

    Returns list of preserved drawing elements (in original order) if preserve_drawings=True.
    Returns anchor element = the paragraph BEFORE start_idx (stable ref for subsequent insertions).
    """
    paras = list(doc.paragraphs)
    anchor = paras[start_idx - 1]._element if start_idx > 0 else None
    to_delete = [paras[i]._element for i in range(start_idx, end_idx + 1)]

    preserved = []
    for el in to_delete:
        if preserve_drawings and has_drawing(el):
            preserved.append(el)
            continue
        el.getparent().remove(el)

    return anchor, preserved


def make_p(doc, text, style='Normal', bold=False, italic=False, size=None,
           align=None, color=None, indent=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p._element


def make_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)
    return p._element


def make_caption(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    return p._element


def make_footnote_para(doc, number, text):
    """Simulates a footnote via italic small paragraph at end (docx footnote API is complex).

    Format: ¹ Text...
    """
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(f'{number} ')
    run.bold = True
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(0.5)
    return p._element


def make_list_item(doc, text, level=0):
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    return p._element


def insert_sequence_after(anchor, elements):
    """Insert elements in order after anchor. Each new element becomes next sibling in sequence."""
    ref = anchor
    for el in elements:
        ref.addnext(el)
        ref = el


def apa_table(doc, headers, rows, title=None, note=None):
    """Create APA-style table: no vertical lines, horizontal lines at top/header/bottom only.

    Returns list of elements: [title_para, table, note_para] (if title/note provided).
    """
    elements = []

    if title:
        t_para = doc.add_paragraph(style='Normal')
        t_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = t_para.add_run(title)
        run.italic = True
        run.font.size = Pt(10)
        elements.append(t_para._element)

    tbl = doc.add_table(rows=1, cols=len(headers))
    # Clear default grid style; apply borders manually
    tbl.style = None

    def set_cell_border(cell, top=False, bottom=False):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)
        # Clear existing
        for side in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
            el = tc_borders.find(qn(f'w:{side}'))
            if el is not None:
                tc_borders.remove(el)
        # No left/right borders (APA)
        for side in ['left', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'nil')
            tc_borders.append(b)
        if top:
            b = OxmlElement('w:top')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:color'), '000000')
            tc_borders.append(b)
        if bottom:
            b = OxmlElement('w:bottom')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:color'), '000000')
            tc_borders.append(b)

    # Header row
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_border(hdr[i], top=True, bottom=True)

    # Data rows
    for r_idx, row in enumerate(rows):
        is_last = (r_idx == len(rows) - 1)
        new_row = tbl.add_row().cells
        for i, v in enumerate(row):
            new_row[i].text = str(v)
            for para in new_row[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < len(headers) - 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
            set_cell_border(new_row[i], bottom=is_last)

    elements.append(tbl._element)

    if note:
        n_para = doc.add_paragraph(style='Normal')
        run = n_para.add_run('Nota. ')
        run.italic = True
        run.font.size = Pt(9)
        run2 = n_para.add_run(note)
        run2.font.size = Pt(9)
        elements.append(n_para._element)

    return elements


# ===== Section content builders =====

def build_justificacion(doc):
    """II.A — Justification with 3 footnote-like items."""
    elems = []
    elems.append(make_p(doc,
        'La realización del presente trabajo de grado surge ante la necesidad de garantizar '
        'la eficiencia y confiabilidad en procesos industriales de taladrado, particularmente '
        'enfocados en la detección temprana del desgaste de herramientas de corte mediante '
        'métodos no invasivos, como el monitoreo acústico. El desgaste de herramientas, '
        'especialmente en procesos de mecanizado, genera pérdidas económicas significativas '
        'debido a tiempos de parada no planificados, reducción de calidad superficial, '
        'incremento en costos de mantenimiento y posibles fallas en componentes críticos del '
        'producto final. Por lo tanto, implementar técnicas que permitan predecir y diagnosticar '
        'el estado de desgaste de herramientas en tiempo real es crucial para optimizar los '
        'procesos industriales, minimizar costos operativos y mantener altos estándares de calidad.'))
    elems.append(make_p(doc,
        'El principal aporte de este trabajo es la construcción de un sistema de monitoreo '
        'acústico multimodal que, mediante la captura y análisis de señales sonoras generadas '
        'durante procesos de taladrado, permite clasificar el grado de desgaste de brocas '
        'utilizando clasificación ordinal¹ basada en máquinas de vectores de soporte (SVM)² y, '
        'como extensión, redes neuronales convolucionales (CNN)³ sobre representaciones '
        'espectrales. El enfoque ordinal —a diferencia de la clasificación binaria empleada en '
        'trabajos previos— preserva el orden natural entre estados de desgaste y garantiza que '
        'los errores de clasificación nunca excedan un paso entre clases adyacentes, propiedad '
        'esencial para un sistema de alerta temprana. Adicionalmente, el sistema integra '
        'variables de apoyo (caudal de refrigerante e inspección visual microscópica) que '
        'enriquecen la base de conocimiento y refuerzan la confiabilidad del diagnóstico.'))
    elems.append(make_p(doc,
        'Este proyecto beneficiará directamente a las industrias manufactureras, especialmente '
        'aquellas dedicadas al mecanizado y producción masiva de componentes metálicos, donde '
        'el desgaste de las herramientas de corte es un factor crítico. El costo del subsistema '
        'auxiliar (ESP32 + micrófono INMP441 + caudalímetro YF-S201) se estima por debajo de '
        '30 USD, lo que convierte la solución en una alternativa viable para pequeñas y '
        'medianas empresas (PYMES) que no disponen de instrumentación de laboratorio costosa. '
        'Además, los resultados y métodos desarrollados pueden ser aprovechados en el ámbito '
        'académico e investigativo, al proporcionar bases para futuras investigaciones '
        'relacionadas con técnicas avanzadas de análisis acústico aplicado a la industria.'))
    # Notes at foot of section (pseudo-footnotes — real footnotes require XML-level editing)
    elems.append(make_footnote_para(doc, '¹',
        'Clasificación ordinal: tipo de clasificación donde las categorías tienen un orden '
        'natural (por ejemplo: sin desgaste < medianamente desgastado < desgastado), a '
        'diferencia de la clasificación categórica donde las clases son independientes entre sí.'))
    elems.append(make_footnote_para(doc, '²',
        'SVM (Máquina de Vectores de Soporte, Support Vector Machine): algoritmo de aprendizaje '
        'automático supervisado que separa clases de datos encontrando la frontera de decisión '
        '(hiperplano) que maximiza la distancia entre las clases más cercanas.'))
    elems.append(make_footnote_para(doc, '³',
        'CNN (Red Neuronal Convolucional, Convolutional Neural Network): arquitectura de '
        'aprendizaje profundo que aprende patrones directamente de imágenes o representaciones '
        'visuales de señales, sin requerir que el investigador defina manualmente qué '
        'características extraer.'))
    return elems


def build_chapter8(doc):
    """II.B — Chapter 8 body rewritten."""
    elems = []
    # Intro (no heading; H1 stays in place as para 791)
    elems.append(make_p(doc,
        'El procesamiento descrito en este capítulo extiende el pipeline acústico del Capítulo 7 '
        'a los ensayos propios del proyecto (Lote II) y lo amplía en dos direcciones: (i) la '
        'fusión multimodal, que incorpora descriptores del caudal de refrigerante y del '
        'recubrimiento de la broca al vector de características acústicas, y (ii) la '
        'clasificación ordinal mediante redes neuronales convolucionales (CNN) sobre '
        'espectrogramas mel, que aprenden representaciones directamente de la señal sin '
        'requerir extracción manual de características. Ambas extensiones se entrenan y evalúan '
        'sobre el conjunto unificado de 4 381 observaciones descrito en el Capítulo 6.'))

    # 8.1
    elems.append(make_heading(doc, '8.1 Segmentación de grabaciones continuas', level=2))
    elems.append(make_p(doc,
        'Las grabaciones multicanal de los ensayos contienen señales acústicas de los eventos '
        'de taladrado, los movimientos rápidos de reposicionamiento, las fases de limpieza '
        'mediante sopladora manual y las pausas de inspección. Para producir segmentos útiles '
        'al modelo de clasificación se desarrolló el algoritmo de segmentación por '
        'CutRatio¹: un cociente espectral entre la energía contenida en la banda alta '
        '(3–12 kHz) y la energía contenida en la banda baja (200–1 500 Hz), calculado por '
        'tramas de 0,5 segundos con solape de 0,25 segundos.'))
    elems.append(make_p(doc,
        'El principio físico subyacente es que durante el corte real el espectro se ensancha '
        'hacia la banda media por la formación de viruta (ruido de banda ancha), mientras que '
        'durante el reposicionamiento entre agujeros el ruido del husillo es tonal y concentra '
        'energía en la banda alta. En consecuencia, el cociente CutRatio aumenta durante los '
        'transitorios de reposicionamiento y disminuye durante el corte sostenido. Este '
        'comportamiento permite delimitar bloques continuos de taladrado y contar los agujeros '
        'individuales.'))
    elems.append(make_p(doc,
        'Tras la segmentación se aplica un margen de 0,5 segundos antes y después de cada pico '
        'detectado para garantizar que la firma acústica completa del corte quede contenida en '
        'el segmento. Los segmentos resultantes se almacenan con nomenclatura estructurada '
        'junto con un manifiesto JSON que registra el tiempo absoluto de inicio, la duración, '
        'la identidad del canal y el hash SHA-256 del archivo.'))
    elems.append(make_footnote_para(doc, '¹',
        'CutRatio: cociente entre la energía espectral de la banda alta y la banda baja de la '
        'señal acústica; valores elevados indican movimiento sin corte, valores bajos indican '
        'taladrado activo.'))

    # 8.2
    elems.append(make_heading(doc, '8.2 Extracción de 26 características acústicas', level=2))
    elems.append(make_p(doc,
        'El extractor de características reutiliza íntegramente el pipeline del Capítulo 7, '
        'con el fin de asegurar la compatibilidad del conjunto de datos ampliado con el modelo '
        'SVM ordinal. Sobre cada segmento de agujero se computan 26 descriptores agrupados en '
        'tres dominios físicos, resumidos en la Tabla 8.1.'))

    t81 = apa_table(doc,
        headers=['Dominio', 'Descriptores', 'Interpretación'],
        rows=[
            ['Energía', 'RMS, RMS en dB, pico, factor de cresta',
             'Potencia, volumen y percusividad de la señal'],
            ['Timbre',
             'Centroide, rolloff, ancho de banda, planitud, entropía, contraste espectral, '
             'MFCC 0–1, chroma, tonnetz',
             'Forma y calidad tonal del espectro'],
            ['Temporal/rítmico',
             'Duración, ZCR, tempo, tasa de onsets, razón armónico-percusivo',
             'Estructura temporal y eventos transitorios'],
        ],
        title='Tabla 8.1. Dominios físicos y descriptores acústicos extraídos por segmento',
        note='RMS = valor cuadrático medio; MFCC = coeficientes cepstrales en escala mel; '
             'ZCR = tasa de cruces por cero. La descripción completa de cada descriptor se '
             'encuentra en el Apéndice A.')
    elems.extend(t81)

    # 8.3
    elems.append(make_heading(doc, '8.3 Descriptores multimodales adicionales', level=2))
    elems.append(make_p(doc,
        'El Artículo 2 de este proyecto (Ayala et al., 2026b) incorpora tres bloques de '
        'descriptores que complementan los 26 descriptores acústicos del sistema NI, '
        'conformando un vector multimodal de 38 dimensiones. La Tabla 8.2 resume la '
        'composición completa.'))

    t82 = apa_table(doc,
        headers=['Bloque', 'Descriptores', 'Cantidad'],
        rows=[
            ['Audio NI', 'Los 26 descriptores de la Tabla 8.1', '26'],
            ['Audio ESP32 (INMP441)',
             'RMS, RMS en dB, centroide, ZCR, contraste espectral, factor de cresta', '6'],
            ['Caudal (YF-S201)',
             'Media, desviación estándar, mínimo, ciclo de trabajo de pulsos, coeficiente de variación', '5'],
            ['Recubrimiento', 'Código binario de presencia de nitruro de titanio (TiN)', '1'],
            ['Total', '', '38'],
        ],
        title='Tabla 8.2. Bloques de descriptores del vector multimodal completo',
        note='Las observaciones del corpus de Orejarena (2014) carecen por construcción de los '
             'canales multimodales (ESP32, caudal y recubrimiento); sus columnas correspondientes '
             'se imputan por la mediana del conjunto de entrenamiento. La cobertura efectiva en '
             'el conjunto unificado es: ESP32 27,3 %, caudal 11,6 %, recubrimiento 33,0 %.')
    elems.extend(t82)

    # 8.4
    elems.append(make_heading(doc, '8.4 Etiquetado ordinal y reentrenamiento incremental', level=2))
    elems.append(make_p(doc,
        'A cada segmento se asigna una etiqueta ordinal basada en la fracción de vida útil '
        'consumida: «sin desgaste» para fracción menor al 15 %, «medianamente desgastado» '
        'entre el 15 % y el 75 %, y «desgastado» por encima del 75 %. El reentrenamiento se '
        'ejecuta automáticamente tras cada ensayo nuevo, orquestando siete etapas secuenciales: '
        'localización del ensayo, segmentación por CutRatio, extracción de las 26 '
        'características, asignación de etiquetas, fusión con el conjunto existente, '
        'reentrenamiento del SVM ordinal y evaluación sobre el conjunto de retención E3. El '
        'ciclo completo se ejecuta en 5 a 7 minutos, lo que permite actualizar el modelo entre '
        'ensayos consecutivos.'))

    # 8.5
    elems.append(make_heading(doc, '8.5 Clasificación ordinal por descomposición de Frank y Hall', level=2))
    elems.append(make_p(doc,
        'La clasificación ordinal aborda un problema cuya variable objetivo tiene un orden '
        'natural («sin desgaste» < «medianamente desgastado» < «desgastado») pero sin '
        'intervalos uniformes entre clases. El método de Frank y Hall (2001) descompone este '
        'problema en dos clasificadores binarios secuenciales:'))
    elems.append(make_list_item(doc,
        'C₁: ¿hay desgaste? (sin desgaste frente a medianamente desgastado + desgastado).'))
    elems.append(make_list_item(doc,
        'C₂: ¿el desgaste es severo? (sin desgaste + medianamente desgastado frente a desgastado).'))
    elems.append(make_p(doc,
        'La probabilidad de la clase intermedia se obtiene por diferencia: P(medianamente) = '
        'P(C₁ = 1) × (1 − P(C₂ = 1)). Este esquema tiene dos ventajas frente a un clasificador '
        'multiclase estándar: conserva el orden en las predicciones —un error de clasificación '
        'tiende a caer en la clase adyacente, no en la extrema opuesta— y permite calibrar '
        'umbrales de decisión independientes para cada clasificador.'))

    # 8.6
    elems.append(make_heading(doc, '8.6 Comparativa SVM tabular vs. CNN sobre espectrogramas mel', level=2))
    elems.append(make_p(doc,
        'Para evaluar si las representaciones aprendidas superan a los descriptores diseñados '
        'manualmente, se entrenaron dos arquitecturas de red neuronal convolucional con la '
        'misma descomposición ordinal de Frank y Hall, sobre espectrogramas mel¹ normalizados '
        'de 64 × 512 píxeles (44,1 kHz, ventana Hanning de 23 ms con solapamiento del 50 %). '
        'El entrenamiento se realizó en una instancia de cómputo en la nube con tarjeta '
        'gráfica (GPU) mediante la plataforma Paperspace, dado que el entrenamiento de redes '
        'profundas requiere capacidad de procesamiento paralelo no disponible en el equipo '
        'local de laboratorio.'))
    elems.append(make_p(doc,
        'CNN-B (red de referencia): capa inicial de convolución seguida de tres bloques que '
        'incrementan los canales de 32 a 256, agrupación promedio global, regularización por '
        'descarte (Dropout 0,4) y una capa de salida de dos neuronas para la descomposición '
        'ordinal. Total: 1,21 millones de parámetros.'))
    elems.append(make_p(doc,
        'CNN-v3 (red profunda): extiende CNN-B con un cuarto bloque de convolución '
        '(256 → 384 canales) y una capa densa intermedia de 64 neuronas. Total: 3,8 millones '
        'de parámetros. A diferencia de CNN-B, el modelo se preserva por el mejor F₁ macro '
        'sobre validación, no por exactitud adyacente, para evitar el colapso hacia la clase '
        'mayoritaria².'))
    elems.append(make_p(doc,
        'Ambas arquitecturas emplean suavizado de etiquetas (ε = 0,05), restricción de '
        'monotonía en inferencia, programación de tasa de aprendizaje OneCycleLR con '
        'η_máx = 3 × 10⁻⁴, 60 épocas, lote de 48 muestras y aumentación por SpecAugment '
        '(enmascarado aleatorio de frecuencia y tiempo más ruido gaussiano). Un ensamblado '
        'pondera los vectores de probabilidad: p̂_ens = 0,2 × p̂_B + 0,8 × p̂_v3. El peso '
        'óptimo se determinó por barrido sobre la retención E3.'))
    elems.append(make_footnote_para(doc, '¹',
        'Espectrograma mel: representación visual de una señal acústica en la que el eje '
        'horizontal es el tiempo, el eje vertical es la frecuencia en escala mel '
        '(perceptualmente uniforme para el oído humano) y el color representa la intensidad. '
        'Permite que la CNN aprenda patrones visuales asociados a cada estado de desgaste.'))
    elems.append(make_footnote_para(doc, '²',
        'Colapso de clase: fenómeno en el que el modelo aprende a predecir siempre la clase '
        'más frecuente (medianamente desgastado), obteniendo exactitud adyacente aparentemente '
        'alta pero con F₁ de la clase desgastado igual a cero.'))

    # 8.7
    elems.append(make_heading(doc, '8.7 Integración multimodal: caudal, video y predicción en tiempo real', level=2))
    elems.append(make_p(doc,
        'El subsistema auxiliar ESP32 añade dos modalidades sensoriales no presentes en el '
        'corpus heredado: un caudalímetro YF-S201 sobre la línea de taladrina, que registra el '
        'flujo volumétrico del refrigerante a 1 Hz mediante conteo de pulsos del sensor de '
        'efecto Hall, y un micrófono digital INMP441 por protocolo I²S. El flujo de '
        'refrigerante se correlaciona con la capacidad térmica instantánea del proceso: '
        'caídas transitorias del caudal indican obstrucción de la línea o sobrecalentamiento, '
        'condiciones que aceleran el desgaste y modifican la firma acústica del corte.'))
    elems.append(make_p(doc,
        'La cámara microscópica CMOS registra el flanco de la broca cada quince agujeros '
        'durante una pausa controlada. Las imágenes se correlacionan temporalmente con los '
        'eventos del registro acústico mediante marcas de tiempo sincronizadas, que mantienen '
        'un desfase menor a 500 milisegundos respecto al arranque de la adquisición NI. La '
        'inspección óptica empleada como sistema de referencia dimensional se detalla en la '
        'sección 8.7.1.'))
    elems.append(make_p(doc,
        'La arquitectura de predicción en tiempo real implementa una ventana deslizante de 20 '
        'segundos con solape, extrae las 26 características en paralelo por canal y emite la '
        'predicción ordinal en menos de un segundo por muestra: una latencia adecuada para '
        'retroalimentación audiovisual al operador durante el ensayo.'))

    # 8.7.1 Mirilla subsection (integrated from previous Ch9 mirilla block)
    elems.append(make_heading(doc, '8.7.1 Medición cuantitativa del desgaste de filo mediante mirilla óptica', level=3))
    elems.append(make_p(doc,
        'Como complemento a la inspección visual cualitativa, se implementó un flujo de '
        'medición cuantitativa del desgaste en el borde del cincel (chisel edge) empleando la '
        'mirilla óptica como sistema de referencia dimensional. La mirilla incorpora una regla '
        'graduada visible en el campo de visión, lo que permite convertir distancias en '
        'píxeles a milímetros mediante una línea de calibración trazada ortogonalmente sobre '
        'la marca de 1 mm de la regla.'))
    elems.append(make_p(doc,
        'El protocolo de medición sobre los fotogramas extraídos consistió en: (1) trazar una '
        'línea de calibración sobre la referencia de 1 mm de la regla para obtener el factor '
        'px/mm de cada sesión de inspección; (2) trazar una línea ortogonal al punto muerto '
        'de la broca sobre la zona de desgaste del chisel edge, desde el filo íntegro hasta '
        'el límite visible del material removido; y (3) convertir la longitud en píxeles al '
        'valor en milímetros usando el factor de calibración del ensayo correspondiente.'))
    elems.append(make_p(doc,
        'La estabilidad del factor de calibración entre ensayos fue alta: los valores '
        'medianos oscilaron entre 162 y 184 px/mm (coeficiente de variación inferior al 8 %), '
        'lo que indica que la posición de la mirilla respecto a la broca se mantuvo '
        'consistente entre sesiones de inspección. En el ensayo 53 (Dormer A114 #1, fractura '
        'en el agujero 44) se obtuvieron cuatro mediciones del chisel edge wear en la misma '
        'parada de inspección, con una media de 0,626 mm y una desviación estándar de 0,009 mm '
        '(CV = 1,4 %), lo que confirma la repetibilidad del método de marcado.'))
    elems.append(make_p(doc,
        'Para el ensayo 39 (Dormer A100, 110 agujeros hasta fractura) se implementó '
        'adicionalmente un algoritmo de detección automática de eventos de inspección basado '
        'en nitidez sostenida (varianza del operador Laplaciano superior a 80 sobre ventanas '
        'de 5 s con salto temporal). El algoritmo identificó 9 eventos de parada distribuidos '
        'a lo largo del ensayo, espaciados en promedio 7,0 minutos —equivalentes a '
        'aproximadamente 15 agujeros por intervalo a la cadencia de taladrado observada '
        '(42 s/agujero)—, lo que valida la correspondencia entre las inspecciones detectadas '
        'automáticamente y el protocolo operativo documentado.'))

    elems.append(make_p(doc,
        'La transición del procesamiento al análisis de resultados se presenta en el siguiente '
        'capítulo, donde se cuantifican los efectos de cada extensión (multimodal y CNN) sobre '
        'las métricas de clasificación.'))

    return elems


def build_chapter9(doc):
    """II.C — Chapter 9 body rewritten (H1 'Resultados y discusión' preserved)."""
    elems = []
    elems.append(make_p(doc,
        'Este capítulo sintetiza los resultados cuantitativos producidos por los pipelines '
        'descritos en los Capítulos 7 y 8. Se organiza en seis bloques: (i) resumen métrico '
        'del modelo SVM acústico, (ii) impacto del reentrenamiento incremental, (iii) fusión '
        'multimodal, (iv) comparativa SVM frente a CNN, (v) rol de las variables de apoyo y '
        '(vi) hallazgos y limitaciones.'))

    # 9.1
    elems.append(make_heading(doc, '9.1 Resumen cuantitativo del modelo SVM acústico', level=2))
    t91 = apa_table(doc,
        headers=['Iteración', 'Muestras', 'Exactitud exacta',
                 'Exactitud adyacente', 'F₁ macro', 'Δ adj. acc.'],
        rows=[
            ['Base — corpus Orejarena (sin spectral gating)', '2 352', '43,2 %', '84,7 %', '0,42', '—'],
            ['Base — corpus Orejarena (+ spectral gating)',    '2 352', '50,3 %', '90,1 %', '0,47', '+5,4 pp'],
            ['Iter. 3 — + ensayo 39',                          '3 271', '51,1 %', '92,3 %', '0,35', '+2,2 pp'],
            ['Iter. 6 — corpus completo (4 381)',              '4 381', '52,0 %', '89,4 %', '0,35', '−2,9 pp'],
            ['Iter. 7 — multimodal completo',                  '4 381', '52,1 %', '95,4 %', '0,35', '+6,0 pp'],
        ],
        title='Tabla 9.1. Evolución de métricas del modelo SVM ordinal Frank–Hall por iteración '
              'de reentrenamiento, evaluado sobre la retención E3 (n = 583)',
        note='La exactitud adyacente considera correcto un error de ±1 clase ordinal. El esquema '
             'de etiquetado [15/75] clasifica el primer 15 % de agujeros como «sin desgaste» y '
             'el último 25 % como «desgastado». Δ adj. acc. = variación respecto a la iteración '
             'anterior. pp = puntos porcentuales.')
    elems.extend(t91)
    elems.append(make_p(doc,
        'El modelo SVM ordinal Frank–Hall, entrenado sobre el corpus unificado con 26 '
        'características acústicas y reducción de ruido espectral (spectral gating, α = 0,85), '
        'alcanza una exactitud adyacente del 90,1 % sobre E3. La incorporación de los ensayos '
        'propios (iteraciones 3 a 6) estabiliza la métrica en torno al 89,4 %, con variaciones '
        'menores a 0,5 puntos porcentuales entre iteraciones consecutivas. La variante '
        'multimodal (iteración 7, 38 descriptores) rompe esta meseta y eleva la exactitud '
        'adyacente a 95,4 % (+6,0 pp), como se analiza en la sección 9.3.'))

    # 9.2
    elems.append(make_heading(doc, '9.2 Impacto del reentrenamiento incremental', level=2))
    elems.append(make_p(doc,
        'La estabilidad de la exactitud adyacente en torno al 89,4 % a lo largo de las '
        'iteraciones 2 a 6 evidencia que la arquitectura SVM Frank–Hall con hiperparámetros '
        'fijos (C = 10, núcleo RBF, 15 características seleccionadas por información mutua) se '
        'encuentra en una meseta de rendimiento. La ampliación del conjunto de datos con '
        '1 446 muestras nuevas aporta diversidad de microfonía y de geometría de broca '
        '(Dormer A100 y A114). El efecto neto es un modelo más robusto al cambio de dominio¹, '
        'a costa de una ligera disminución del ajuste sobre E3 —disminución cuantitativamente '
        'irrelevante frente al beneficio de generalización.'))
    elems.append(make_footnote_para(doc, '¹',
        'Cambio de dominio (domain shift): diferencia estadística entre los datos de '
        'entrenamiento y los de prueba, causada aquí por variaciones en el tipo de micrófono, '
        'la geometría de la broca o las condiciones ambientales entre lotes experimentales.'))

    # 9.3
    elems.append(make_heading(doc, '9.3 Fusión multimodal: ablación de variantes', level=2))
    elems.append(make_p(doc,
        'La Tabla 9.2 compara tres variantes del clasificador SVM ordinal sobre la retención '
        'E3, lo que permite evaluar la contribución individual de cada bloque de descriptores.'))
    t92 = apa_table(doc,
        headers=['Variante', 'Descriptores', 'Exactitud exacta',
                 'Exactitud adyacente', 'F₁ macro', 'MAE ordinal'],
        rows=[
            ['A. Solo audio',              '26', '0,520', '0,894', '0,351', '0,587'],
            ['B. Audio + recubrimiento',   '27', '0,509', '0,885', '0,347', '0,605'],
            ['C. Multimodal completo',     '38', '0,521', '0,954', '0,354', '0,525'],
        ],
        title='Tabla 9.2. Desempeño comparativo de las variantes del clasificador SVM sobre E3 (n = 583)',
        note='MAE ordinal = error absoluto medio entre clases predichas y verdaderas. La variante '
             'B no incluye descriptores de caudal; la variante C incluye audio NI, audio ESP32, '
             'caudal y recubrimiento.')
    elems.extend(t92)
    elems.append(make_p(doc,
        'La variante B (audio + recubrimiento aislado) no mejora al modelo base e incluso lo '
        'degrada en −0,9 pp de exactitud adyacente. El recubrimiento, como variable aislada, '
        'queda confundido con la identidad de la broca y no aporta capacidad discriminativa. '
        'Sin embargo, al combinarse con el bloque hidráulico (variante C), la variable '
        'categórica adquiere valor porque el caudal traduce el estado físico de la broca a una '
        'señal continua correlacionada con el fenómeno de desgaste. Este hallazgo reafirma un '
        'principio general de fusión multimodal: una covariable categórica aporta poco si no '
        'está anclada a una covariable continua correlacionada con el fenómeno físico '
        'subyacente.'))
    elems.append(make_p(doc,
        'La ganancia de 6,0 pp en exactitud adyacente proviene casi íntegramente de la '
        'reducción de errores a dos pasos ordinales en la clase «desgastado»: los casos de '
        'verdadero desgastado clasificados como sin desgaste caen de 62 a 27, una disminución '
        'del 56,5 %.'))

    # 9.4
    elems.append(make_heading(doc, '9.4 Comparativa SVM vs. CNN ordinal', level=2))
    elems.append(make_p(doc,
        'Para superar la meseta del modelo tabular, se entrenaron las arquitecturas CNN '
        'descritas en la sección 8.6. La Tabla 9.3 compara cinco configuraciones sobre la '
        'retención E3.'))
    t93 = apa_table(doc,
        headers=['Modelo', 'Exactitud exacta', 'Exactitud adyacente', 'F₁ macro', 'F₁,deg'],
        rows=[
            ['SVM acústico (Art. 1)',       '0,520', '0,894', '0,351', '0,328'],
            ['SVM multimodal (Art. 2)',     '0,521', '0,954', '0,354', '0,340'],
            ['CNN-B',                       '0,566', '0,988', '0,445', '0,403'],
            ['CNN-v3',                      '0,566', '0,959', '0,529', '0,545'],
            ['Ensamblado (0,2B + 0,8v3)',   '0,576', '0,967', '0,531', '0,552'],
        ],
        title='Tabla 9.3. Comparativa de modelos sobre E3 (n = 583)',
        note='F₁,deg = F₁ de la clase desgaste severo. El SVM acústico corresponde al modelo '
             'del Artículo 1 (Ayala et al., 2026a); el SVM multimodal corresponde a la '
             'variante C del Artículo 2 (Ayala et al., 2026b). Las CNN se entrenaron en una '
             'instancia con GPU en Paperspace.')
    elems.extend(t93)
    elems.append(make_p(doc,
        'CNN-B logra la exactitud adyacente más alta (98,8 %, +8,7 pp sobre SVM Art. 1), lo '
        'que evidencia que el espectrograma mel captura patrones tiempo-frecuencia no '
        'accesibles para los 26 descriptores estadísticos. El ensamblado CNN-B/v3 maximiza '
        'F₁,deg = 0,552, un incremento del 62 % relativo sobre el SVM multimodal (0,340) para '
        'la clase crítica de seguridad. Esta métrica es la más relevante operativamente: '
        'detectar correctamente el desgaste severo es prioritario porque un falso negativo de '
        'dos pasos («sin desgaste» cuando la broca está desgastada) puede provocar la '
        'fractura de la herramienta dentro de la pieza.'))
    elems.append(make_p(doc,
        'Sin embargo, el SVM multimodal conserva ventajas operativas significativas. Primero, '
        'el análisis SHAP permite explicar cada predicción a nivel de descriptor físico, lo '
        'que facilita la auditoría y la detección de fallos de sensor en campo. Segundo, su '
        'costo computacional es mínimo (inferencia inferior a 1 ms por muestra), no requiere '
        'GPU y puede ejecutarse en sistemas embebidos. Para despliegue en PYMES donde el '
        'acceso a expertos en aprendizaje profundo es limitado, la variante SVM multimodal '
        'sigue siendo competitiva (adj. acc. 95,4 %, sólo −3,4 pp bajo la mejor CNN) con '
        'máxima interpretabilidad.'))

    # 9.5
    elems.append(make_heading(doc, '9.5 Rol de las variables de apoyo: caudal y video microscópico', level=2))
    elems.append(make_p(doc,
        'El análisis multimodal sobre el ensayo 53 (Dormer A114 #1, fractura en el agujero 44) '
        'muestra una correlación cualitativa clara entre tres indicadores: la probabilidad de '
        'desgaste estimada por el modelo exhibe picos crecientes a partir del agujero 30 en '
        'los tres micrófonos, el caudal de taladrina presenta pausas que coinciden con las '
        'inspecciones visuales, y los fotogramas microscópicos documentan la progresión del '
        'flanco desde filo íntegro hasta fractura final.'))
    elems.append(make_p(doc,
        'El caudal medio (flow_mean) exhibe correlaciones negativas crecientes en magnitud '
        'conforme avanza el desgaste (Spearman ρ = −0,42 en sin desgaste; −0,76 en moderado; '
        '−0,81 en severo). El coeficiente de variación del caudal invierte su signo entre sin '
        'desgaste (+0,71) y severo (−0,12), reflejando la regularización del régimen hidráulico '
        'al degradarse la broca. Esta inversión de signo constituye el marcador diagnóstico '
        'más claro de la transición hacia el estado de desgaste severo.'))
    elems.append(make_p(doc,
        'Físicamente, este patrón es consistente con el modelo tribológico de Shaw (2005): '
        'una broca desgastada presenta mayor resistencia al corte, lo que aumenta la fricción '
        'térmica en la interfaz herramienta-viruta y reduce las fluctuaciones de contrapresión '
        'en el circuito hidráulico. En términos prácticos, el caudalímetro YF-S201 actúa como '
        'un indicador indirecto del estado térmico en la zona de corte, que de otro modo '
        'requeriría instrumentación intrusiva.'))

    # 9.6
    elems.append(make_heading(doc, '9.6 Hallazgos y limitaciones', level=2))
    elems.append(make_p(doc,
        'Hallazgo central: la propiedad «errores acotados a una clase ordinal» persiste en '
        'todos los modelos evaluados, tanto sobre el corpus inicial como sobre el corpus '
        'ampliado. Esta propiedad tiene valor práctico para el mantenimiento predictivo: un '
        'sistema que nunca confunde «sin desgaste» con «desgastado» puede tolerarse en un '
        'esquema de alerta temprana con umbrales conservadores, aun cuando la exactitud '
        'exacta sea moderada.'))
    elems.append(make_p(doc, 'Limitaciones identificadas:'))
    elems.append(make_list_item(doc,
        '(1) La cobertura efectiva del caudal en el conjunto unificado es de sólo 11,6 %. La '
        'imputación por mediana puede ocultar parte del efecto real; ampliar la proporción de '
        'ensayos con instrumentación completa debería ser prioritario.'))
    elems.append(make_list_item(doc,
        '(2) La base de Orejarena no permite desacoplar el efecto del micrófono del efecto del '
        'experimento, lo que introduce un factor de confusión.'))
    elems.append(make_list_item(doc,
        '(3) La clase «desgastado» está subrepresentada por construcción (25 % por broca), y '
        'ningún modelo SVM recupera aciertos exactos sobre E3 para esta clase.'))
    elems.append(make_list_item(doc,
        '(4) La validación fuera de dominio (LODO) se ejecutó sobre una sola broca A114. '
        'Replicar sobre las brocas #2 y #3 reforzaría la evidencia.'))
    elems.append(make_list_item(doc,
        '(5) Se observó una degradación de −10,9 pp sobre E2 al usar la variante multimodal. '
        'E2 pertenece a Orejarena y carece totalmente de cobertura multimodal; la imputación '
        'por mediana introduce un sesgo sistemático. Esto constituye una advertencia '
        'metodológica: la imputación por mediana es adecuada cuando la ausencia del canal es '
        'aleatoria, pero no cuando es estructural.'))
    elems.append(make_p(doc,
        'La síntesis de cumplimiento de objetivos, contribuciones y trabajo futuro se presenta '
        'en el capítulo siguiente.'))

    return elems


def build_chapter10(doc):
    """II.D — Chapter 10 body rewritten."""
    elems = []
    elems.append(make_p(doc,
        'El presente trabajo de grado construye una base de conocimiento operativa para el '
        'monitoreo acústico del desgaste de herramientas en taladrado CNC, materializada en: '
        'un sistema de adquisición multimodal (audio NI de alta fidelidad + caudal ESP32 + '
        'inspección visual microscópica), un pipeline reproducible de procesamiento y '
        'clasificación, y un conjunto de datos de 4 381 muestras etiquetadas ordinalmente.'))

    # 10.1
    elems.append(make_heading(doc, '10.1 Cumplimiento de objetivos', level=2))
    elems.append(make_p(doc,
        'Objetivo general — desarrollo de una base de conocimiento para la clasificación del '
        'desgaste de broca. Se considera satisfecho: el sistema integrado adquiere, segmenta, '
        'clasifica y visualiza el estado de la herramienta con latencia inferior a un segundo '
        'por ventana.'))
    elems.append(make_p(doc,
        'OE1 — Rediseño experimental: cumplido mediante el protocolo documentado en el '
        'Capítulo 4, con nueve ensayos propios que complementan los siete ensayos heredados.'))
    elems.append(make_p(doc,
        'OE2 — Sistema de adquisición: cumplido mediante el hardware NI cDAQ-9174 + NI-9234, '
        'el subsistema ESP32 + INMP441 + YF-S201 y la interfaz gráfica desarrollada en Python '
        '(Capítulo 5).'))
    elems.append(make_p(doc,
        'OE3 — Clasificador acústico: cumplido. Se implementaron dos familias de '
        'clasificadores ordinales con descomposición de Frank y Hall: un SVM sobre 26 '
        'características diseñadas manualmente (exactitud adyacente 90,1 % en E3) y dos '
        'arquitecturas CNN sobre espectrogramas mel (exactitud adyacente 98,8 % para CNN-B). '
        'Estos resultados se documentan en los artículos científicos asociados (Ayala et al., '
        '2026a; 2026b).'))
    elems.append(make_p(doc,
        'OE4 — Base de conocimiento etiquetada: cumplido. El conjunto unificado de 4 381 '
        'muestras con etiquetado ordinal en tres niveles, trazabilidad por broca, micrófono y '
        'experimento, y metadatos completos, está disponible para investigaciones futuras.'))

    # 10.2
    elems.append(make_heading(doc, '10.2 Contribuciones', level=2))
    elems.append(make_p(doc, 'Contribuciones científicas:'))
    elems.append(make_list_item(doc,
        '(1) Demostración empírica de que la exactitud adyacente en clasificación ordinal de '
        'desgaste se mantiene estable (> 83 %) en un rango amplio de umbrales de etiquetado '
        '(60 %–97 %), lo que permite recalibrar modelos sin reentrenamiento completo.'))
    elems.append(make_list_item(doc,
        '(2) Evidencia de que la fusión multimodal (audio + caudal + recubrimiento) mejora la '
        'clasificación ordinal en +6,0 pp de exactitud adyacente sobre la retención histórica, '
        'y que el recubrimiento sólo aporta cuando se combina con el bloque hidráulico.'))
    elems.append(make_list_item(doc,
        '(3) Las arquitecturas CNN ordinales sobre espectrogramas mel superan al SVM en todas '
        'las métricas: CNN-B alcanza exactitud adyacente del 98,8 % (+8,7 pp sobre el SVM '
        'acústico) y el ensamblado CNN-B/v3 maximiza F₁,deg = 0,552 (+62 % relativo sobre el '
        'SVM).'))
    elems.append(make_list_item(doc,
        '(4) Identificación del contraste espectral medio como característica dominante para '
        'desgaste severo, y de la ventaja del subconjunto top-7 por información mutua.'))
    elems.append(make_list_item(doc,
        '(5) Caracterización del caudal como indicador indirecto del estado térmico: la '
        'inversión del coeficiente de variación del caudal entre sin desgaste (+0,71) y severo '
        '(−0,12) constituye el marcador diagnóstico del desgaste avanzado.'))
    elems.append(make_p(doc, 'Contribuciones prácticas:'))
    elems.append(make_list_item(doc,
        '(1) Conjunto de datos multimodal etiquetado compatible con scikit-learn y '
        'PyTorch/TensorFlow, disponible con DOI persistente.'))
    elems.append(make_list_item(doc,
        '(2) El costo del subsistema auxiliar (ESP32 + INMP441 + YF-S201) inferior a 30 USD '
        'hace viable la replicación en entornos PYME.'))
    elems.append(make_list_item(doc,
        '(3) El SVM multimodal preserva exactitud adyacente del 95,4 % con mínima '
        'infraestructura computacional (inferencia inferior a 1 ms, sin GPU), siendo la opción '
        'recomendada para entornos sin soporte de aprendizaje profundo.'))

    # 10.3
    elems.append(make_heading(doc, '10.3 Recomendaciones y trabajo futuro', level=2))
    elems.append(make_list_item(doc,
        '(1) Estandarizar la microfonía en campañas futuras: un par único dinámico + '
        'condensador común a todos los ensayos.'))
    elems.append(make_list_item(doc,
        '(2) Fijar IP estática en el firmware del ESP32 para evitar fallas por reasignación DHCP.'))
    elems.append(make_list_item(doc,
        '(3) Ejecutar la prueba definitiva de seis discos con la configuración óptima como '
        'lote final de validación.'))
    elems.append(make_list_item(doc,
        '(4) Explorar adaptación de dominio explícita para migrar el modelo al subsistema MCU '
        'sin pérdida significativa de exactitud.'))
    elems.append(make_list_item(doc,
        '(5) Publicar el conjunto de datos curado con DOI persistente siguiendo el estándar de '
        'tarjeta de datos (Gebru et al., 2021).'))
    elems.append(make_list_item(doc,
        '(6) Replicar la validación LODO sobre las brocas A114 #2 y #3 para reforzar la '
        'evidencia de transferencia fuera de dominio.'))
    elems.append(make_list_item(doc,
        '(7) Evaluar imputación condicional por grupo o enmascaramiento del canal multimodal '
        'en inferencia como alternativa a la imputación por mediana.'))

    return elems


def build_alternativa_b(doc):
    """III.4 — Rewrite Alternativa B in third-person impersonal."""
    elems = []
    elems.append(make_p(doc,
        'La Alternativa B propone desacoplar la instrumentación del chasis de la máquina y '
        'ubicarla sobre soportes independientes. En lugar de fijar micrófonos o electrónica al '
        'armazón del taladro CNC, se plantea emplear un par de caballetes robustos apoyados '
        'directamente en el suelo, ubicados en posición lateral respecto a la pieza de trabajo. '
        'Esta configuración persigue minimizar la transmisión directa de vibraciones y '
        'facilitar el ajuste fino de la posición y orientación de cada micrófono.'))
    elems.append(make_p(doc,
        'Los caballetes se conciben como estructuras sólidas de perfiles de acero '
        '(tubo rectangular o perfil en U) con refuerzos triangulados para asegurar rigidez '
        'frente a flexiones transversales. Sobre la cara superior de cada caballete se dispone '
        'una placa de montaje con anclajes normalizados para los soportes de micrófono y para '
        'los módulos de adquisición auxiliares. El desacoplamiento mecánico respecto a la '
        'máquina se completa mediante amortiguadores elastoméricos en la base de cada '
        'caballete.'))
    elems.append(make_p(doc,
        'La separación de la masa instrumental de la masa de la máquina reduce la '
        'transferencia directa de vibraciones al sistema de micrófonos; la operación del '
        'husillo y el contacto herramienta-pieza continúan generando señal acústica '
        'propagada por el aire, que constituye la variable de interés del experimento. En '
        'contrapartida, esta configuración implica mayor ocupación de espacio en el entorno '
        'del banco y requiere verificación de modos propios de vibración de los caballetes '
        'para evitar resonancias que contaminen la señal.'))
    elems.append(make_p(doc,
        'En síntesis, la Alternativa B maximiza el aislamiento vibratorio a costa de mayor '
        'ocupación de espacio y de un análisis modal previo de la estructura de soporte. La '
        'viabilidad final depende de la disponibilidad de área útil en el laboratorio y de la '
        'tolerancia del procedimiento experimental a la posible introducción de resonancias '
        'estructurales por los propios caballetes.'))
    return elems


def build_new_references(doc):
    """III.6 — New references to append at end of references section."""
    elems = []
    elems.append(make_p(doc,
        'Ayala, O. I., Orejarena, N., & Meneses, J. E. (2026a). Clasificación ordinal del '
        'desgaste de brocas mediante análisis acústico multicanal y descomposición de '
        'Frank–Hall. Revista Colombiana de Tecnologías de Avanzada (RCTA). En revisión.'))
    elems.append(make_p(doc,
        'Ayala, O. I., Orejarena, N., & Meneses, J. E. (2026b). Fusión multimodal '
        'acústico-hidráulica para la detección temprana del desgaste de herramientas de '
        'corte en taladrado CNC. Revista Colombiana de Tecnologías de Avanzada (RCTA). '
        'En revisión.'))
    elems.append(make_p(doc,
        'Shaw, M. C. (2005). Metal Cutting Principles (2ª ed.). Oxford University Press.'))
    elems.append(make_p(doc,
        'Frank, E., & Hall, M. (2001). A simple approach to ordinal classification. En '
        'L. De Raedt & P. Flach (Eds.), Machine Learning: ECML 2001 (pp. 145–156). Springer.'))
    elems.append(make_p(doc,
        'Gebru, T., Morgenstern, J., Vecchione, B., Vaughan, J. W., Wallach, H., Daumé III, H., '
        '& Crawford, K. (2021). Datasheets for datasets. Communications of the ACM, 64(12), '
        '86–92. https://doi.org/10.1145/3458723'))
    return elems


# ===== Dedicatoria / agradecimientos fixes =====

def fix_text_in_paragraph(p, replacements):
    """Apply literal string replacements in a paragraph's runs (preserving as much as possible)."""
    for run in p.runs:
        for old, new in replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)


# ===== Main =====

def main():
    print(f'Loading {SRC}')
    doc = Document(str(SRC))

    # ---- Capture ALL anchors upfront BEFORE any deletion (so indices stay valid) ----
    paras = list(doc.paragraphs)

    anchors = {}
    ranges = {}

    # Justification
    anchors['justif'] = paras[IDX['justif_head']]._element
    ranges['justif'] = [paras[i]._element for i in range(IDX['justif_p1'], IDX['justif_p3'] + 1)]

    # Alternativa B
    anchors['altB'] = paras[IDX['altB_head']]._element
    # Keep drawing 467 (figure), delete surrounding text 465-475 except drawing
    ranges['altB_before'] = [paras[i]._element for i in (465,)]  # first para just after heading
    ranges['altB_after'] = [paras[i]._element for i in (468, 469, 470, 471, 472, 474, 475)]  # skip drawing 467 and caption 466

    # Chapter 8 (791 is heading, 792-821 are body — preserve drawings)
    anchors['ch8'] = paras[IDX['ch8_head']]._element
    ch8_to_delete = []
    for i in range(IDX['ch8_head'] + 1, IDX['ch8_end'] + 1):
        el = paras[i]._element
        if not has_drawing(el):
            ch8_to_delete.append(el)
    ranges['ch8'] = ch8_to_delete

    # Chapter 9 (822 heading, 823-869 body — preserve drawings)
    anchors['ch9'] = paras[IDX['ch9_head']]._element
    ch9_to_delete = []
    for i in range(IDX['ch9_head'] + 1, IDX['ch9_end'] + 1):
        el = paras[i]._element
        if not has_drawing(el):
            ch9_to_delete.append(el)
    ranges['ch9'] = ch9_to_delete

    # Chapter 10 (870 heading, 871-895 body)
    anchors['ch10'] = paras[IDX['ch10_head']]._element
    ranges['ch10'] = [paras[i]._element for i in range(IDX['ch10_head'] + 1, IDX['ch10_end'] + 1)]

    # Orphan anexos (896-918)
    ranges['orphan_anexos'] = [paras[i]._element for i in range(IDX['orphan_anexos_start'],
                                                                   IDX['orphan_anexos_end'] + 1)]

    # Orphan tail (988-1058)
    ranges['orphan_tail'] = [paras[i]._element for i in range(IDX['orphan_tail_start'],
                                                                  IDX['orphan_tail_end'] + 1)]

    # References anchor for adding new refs
    anchors['refs_head'] = paras[919]._element

    # ---- Apply deletions (bottom-up for safety even though anchor captured) ----
    print('Deleting orphan tail (988-1058)...')
    for el in ranges['orphan_tail']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting orphan anexos (896-918)...')
    for el in ranges['orphan_anexos']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting chapter 10 body (871-895)...')
    for el in ranges['ch10']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting chapter 9 body (preserving drawings)...')
    for el in ranges['ch9']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting chapter 8 body (preserving drawings)...')
    for el in ranges['ch8']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting Alternativa B text (preserving drawing+caption)...')
    for el in ranges['altB_before'] + ranges['altB_after']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    print('Deleting Justification body...')
    for el in ranges['justif']:
        if el.getparent() is not None:
            el.getparent().remove(el)

    # ---- Build and insert new sections (anchors captured before deletion, still valid XML refs) ----
    print('Inserting new Justification...')
    new_justif = build_justificacion(doc)
    insert_sequence_after(anchors['justif'], new_justif)

    print('Inserting new Alternativa B body...')
    new_altB = build_alternativa_b(doc)
    insert_sequence_after(anchors['altB'], new_altB)

    print('Inserting new chapter 8 body...')
    new_ch8 = build_chapter8(doc)
    insert_sequence_after(anchors['ch8'], new_ch8)

    print('Inserting new chapter 9 body...')
    new_ch9 = build_chapter9(doc)
    insert_sequence_after(anchors['ch9'], new_ch9)

    print('Inserting new chapter 10 body...')
    new_ch10 = build_chapter10(doc)
    insert_sequence_after(anchors['ch10'], new_ch10)

    print('Appending new references...')
    new_refs = build_new_references(doc)
    insert_sequence_after(anchors['refs_head'], new_refs)

    # ---- Dedicatoria/agradecimientos typo fixes ----
    print('Fixing dedicatoria/agradecimientos typos...')
    # Re-fetch paragraphs (indices may have shifted, but we fix by scanning text patterns)
    for p in doc.paragraphs:
        txt_full = ''.join(r.text for r in p.runs)
        if 'ensenarme' in txt_full or 'aportados' in txt_full or 'peso en oro esta mujer' in txt_full or 'sociedades sanas ha sido la que tuvo' in txt_full:
            fix_text_in_paragraph(p, [
                ('ensenarme', 'enseñarme'),
                ('aportados', 'aportado'),
                ('por tal motivo realmente vale su peso en oro esta mujer',
                 'por lo cual guardo hacia ella un profundo agradecimiento y admiración'),
                ('También ha sido un gran ejemplo sobre las sociedades sanas ha sido la que tuvo',
                 'También ha sido un gran ejemplo de las sociedades sanas que formé'),
            ])

    # ---- Save ----
    print(f'Saving {OUT}')
    doc.save(str(OUT))
    print(f'Total paragraphs in v7: {len(doc.paragraphs)}')
    print(f'Total tables in v7: {len(doc.tables)}')


if __name__ == '__main__':
    main()
