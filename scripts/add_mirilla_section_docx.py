#!/usr/bin/env python3
"""
Agrega sección de medición cuantitativa de chisel edge wear mediante mirilla
al Word Informe del proyecto_v3.docx, después del párrafo 822.
Crea copia: Informe del proyecto_v4.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import copy

SRC  = Path('C:/Users/ayala/Downloads/Informe del proyecto_v3.docx')
OUT  = Path('C:/Users/ayala/Downloads/Informe del proyecto_v4.docx')
FIG  = Path('D:/pipeline_SVM/informe_proyecto/figures/fig_cap9_mirilla.png')

INSERT_AFTER = 822   # insertar después de este índice de párrafo


def add_paragraph_after(doc, ref_para, text, style='Normal', bold=False,
                         italic=False, size=None, alignment=None):
    """Inserta un nuevo párrafo después de ref_para."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    new_para = doc.add_paragraph(style=style)
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = italic
    if size:
        run.font.size = Pt(size)
    if alignment:
        new_para.alignment = alignment

    # Mover el párrafo al lugar correcto (justo después de ref_para)
    ref_para._element.addnext(new_para._element)
    return new_para


def add_picture_after(doc, ref_para, img_path, width_inches=5.5):
    """Inserta un párrafo con imagen centrada después de ref_para."""
    new_para = doc.add_paragraph()
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))
    ref_para._element.addnext(new_para._element)
    return new_para


def main():
    doc = Document(str(SRC))
    paras = doc.paragraphs

    ref = paras[INSERT_AFTER]

    # ── Bloque a insertar (en orden inverso porque cada vez se inserta *después* de ref) ──
    # El último en código queda al principio del bloque por ser insertado primero sobre ref.
    # Para mantener orden: insertar desde el último al primero.

    contenido = [
        # (tipo, texto)
        ('heading2', 'Medición cuantitativa del desgaste de filo mediante mirilla óptica'),
        ('normal',
         'Como complemento a la inspección visual cualitativa descrita anteriormente, se '
         'implementó un flujo de medición cuantitativa del desgaste en el borde del cincel '
         '(chisel edge) empleando la mirilla óptica como sistema de referencia dimensional. '
         'La mirilla incorpora una regla de escala graduada visible en el campo de visión '
         '(Figura 8.2), lo que permite convertir distancias en píxeles a milímetros mediante '
         'una línea de calibración trazada ortogonalmente sobre la marca de 1 mm de la regla.'),
        ('placeholder',
         '[INSERTAR AQUÍ: imagen de la mirilla con la regla en mm señalada — '
         'Figura 8.2. Mirilla óptica empleada para la inspección del filo. '
         'Se indica la regla de referencia graduada en milímetros visible en el campo '
         'de visión de la cámara microscópica CMOS.]'),
        ('normal',
         'El protocolo de medición sobre los fotogramas extraídos consistió en: '
         '(1) trazar una línea de calibración sobre la referencia de 1 mm de la regla para '
         'obtener el factor px/mm de cada sesión de inspección; '
         '(2) trazar una línea ortogonal al punto muerto de la broca sobre la zona de '
         'desgaste del chisel edge, desde el filo íntegro hasta el límite visible del '
         'material removido; y '
         '(3) convertir la longitud en píxeles al valor en milímetros usando el factor '
         'de calibración del ensayo correspondiente.'),
        ('normal',
         'La estabilidad del factor de calibración entre ensayos fue alta: los valores '
         'medianos oscilaron entre 162 y 184 px/mm (coeficiente de variación < 8 %), lo que '
         'indica que la posición de la mirilla respecto a la broca se mantuvo consistente '
         'entre sesiones de inspección. En el ensayo 53 (Dormer A114 #1, fractura en el '
         'agujero 44) se obtuvieron cuatro mediciones del chisel edge wear en la misma '
         'parada de inspección, con una media de 0,626 mm y una desviación estándar de '
         '0,009 mm (CV = 1,4 %), confirmando la repetibilidad del método de marcado.'),
        ('normal',
         'Para el ensayo 39 (Dormer A100, 110 agujeros hasta fractura) se implementó '
         'adicionalmente un algoritmo de detección automática de eventos de inspección '
         'basado en nitidez sostenida (varianza del operador Laplaciano > 80 sobre ventanas '
         'de 5 s con seek temporal). El algoritmo identificó 9 eventos de parada distribuidos '
         'a lo largo del ensayo, espaciados en promedio 7,0 minutos —equivalentes a '
         'aproximadamente 15 agujeros por intervalo a la cadencia de taladrado observada '
         '(42 s/agujero)—, validando la correspondencia entre las inspecciones detectadas '
         'automáticamente y el protocolo operativo documentado.'),
        ('normal',
         'La Figura 8.3 presenta: (A) un fotograma del ensayo 53 con la línea de '
         'calibración (cian) y la medida del chisel edge (rojo, 0,63 mm) anotadas; '
         '(B–D) tres fotogramas del ensayo 39 correspondientes a los eventos de '
         'inspección temprano (h≈15), medio (h≈45) y tardío (h≈93), que ilustran '
         'cualitativamente la progresión del desgaste a lo largo de la vida útil de la '
         'herramienta; (E–H) fotogramas representativos de otros ensayos obtenidos por '
         'detección automática.'),
        ('caption',
         'Figura 8.3. Inspección visual cuantitativa del filo mediante mirilla óptica. '
         '(A) Medición del chisel edge wear en ensayo 53: línea cian = referencia 1 mm, '
         'línea roja = desgaste medido = 0,63 mm. (B–D) Progresión visual en ensayo 39 '
         'desde h≈15 hasta h≈93 agujeros. (E–H) Detección automática en otros ensayos. '
         'Factor de calibración mediano: 164 px/mm (CV < 8 %).'),
    ]

    # Insertar imagen primero (queda última porque se inserta después de ref)
    # luego párrafos en orden inverso para que queden en el orden deseado.

    # 1. Insertar caption (último elemento, insertado después de ref → se moverá al final)
    # Estrategia: insertar en orden INVERSO
    # Estrategia: siempre insertar addnext sobre el nodo FIJO ref._element
    # → el último en insertarse queda primero (orden inverso de inserción)
    # Orden deseado después de ref(822): heading, p1, p2, p3, p4, figura, caption
    # → insertar en orden: caption, figura, p4, p3, p2, p1, heading

    def make_para(doc, tipo, texto):
        if tipo == 'heading2':
            p = doc.add_paragraph(style='Heading 2')
            p.add_run(texto)
        elif tipo == 'caption':
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(texto)
            run.italic = True
            run.font.size = Pt(9)
        elif tipo == 'placeholder':
            p = doc.add_paragraph(style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(texto)
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
            run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(style='Normal')
            p.add_run(texto)
        return p

    # Insertar siempre addnext sobre ref fijo → orden inverso de inserción = orden correcto en doc
    # Orden deseado: heading, p_intro, placeholder_img, p_protocol, p_stability, p_test39, p_fig_desc, fig, caption
    # Insertar: caption, fig, p_fig_desc, p_test39, p_stability, p_protocol, placeholder_img, p_intro, heading

    # caption (queda último)
    cap_para = make_para(doc, 'caption', contenido[-1][1])
    ref._element.addnext(cap_para._element)

    # figura (va antes de caption)
    if FIG.exists():
        fig_para = doc.add_paragraph()
        fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fig_para.add_run()
        r.add_picture(str(FIG), width=Inches(5.8))
        ref._element.addnext(fig_para._element)
    else:
        print(f'WARN: figura no encontrada: {FIG}')

    # párrafos de texto en orden inverso (índice 5 → 0)
    for tipo, texto in reversed(contenido[:-1]):
        p = make_para(doc, tipo, texto)
        ref._element.addnext(p._element)

    doc.save(str(OUT))
    print(f'Guardado: {OUT}')


if __name__ == '__main__':
    main()
