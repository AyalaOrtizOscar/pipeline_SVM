#!/usr/bin/env python3
"""Inspecciona estructura de Informe_del_proyecto_v6.docx."""
from docx import Document
from pathlib import Path
import json

SRC = Path('C:/Users/ayala/Downloads/files/Informe_del_proyecto_v6.docx')
OUT_MAP = Path('D:/pipeline_SVM/article1/v6_structure.json')

doc = Document(str(SRC))

paras = []
for idx, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ''
    txt = p.text.strip()
    if txt or 'Heading' in style:
        paras.append({
            'idx': idx,
            'style': style,
            'text': txt[:180]
        })

OUT_MAP.write_text(json.dumps(paras, ensure_ascii=False, indent=2), encoding='utf-8')

print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Non-empty paragraphs: {len(paras)}')
print(f'Total tables: {len(doc.tables)}')
print(f'Sections: {len(doc.sections)}')
print(f'Map written to {OUT_MAP}')

print('\n--- Headings ---')
for p in paras:
    if 'Heading' in p['style'] or 'Título' in p['style'] or 'Titulo' in p['style']:
        print(f"  [{p['idx']:4d}] {p['style']:20s} | {p['text'][:120]}")
