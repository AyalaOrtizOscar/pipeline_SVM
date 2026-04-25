#!/usr/bin/env python3
"""Sanity-check content of key v7 sections."""
from docx import Document
from pathlib import Path
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = Path('C:/Users/ayala/Downloads/files/Informe_del_proyecto_v7.docx')
doc = Document(str(SRC))

print('=== Justificación (242-250) ===')
for i, p in enumerate(doc.paragraphs):
    if 242 <= i <= 250:
        t = p.text.strip()[:200]
        if t:
            print(f'  [{i:4d}] {t}')

print('\n=== Check for known problematic strings ===')
keywords = {
    'LSTM': 0,
    'K-Means': 0,
    'pipe_clf': 0,
    'compute_thresholds': 0,
    'Imagina sacar': 0,
    'bala de plata': 0,
    'ensenarme': 0,
    'Ayala, O. I., Orejarena': 0,
    'Shaw, M. C.': 0,
    'Paperspace': 0,
    'CNN-B': 0,
    'CNN-v3': 0,
    'Tabla 9.2': 0,
    'Tabla 9.3': 0,
}
for p in doc.paragraphs:
    t = p.text
    for kw in keywords:
        if kw in t:
            keywords[kw] += 1
# Also check tables
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for kw in keywords:
                if kw in cell.text:
                    keywords[kw] += 1

for kw, count in keywords.items():
    marker = '✗' if ((kw in ['LSTM', 'K-Means', 'pipe_clf', 'compute_thresholds',
                             'Imagina sacar', 'bala de plata', 'ensenarme']) and count > 0) else '✓'
    print(f'  {marker} "{kw}": {count}')

print('\n=== Dedicatoria typos ===')
for i, p in enumerate(doc.paragraphs):
    if i < 65:
        t = p.text
        if 'ensenar' in t.lower() or 'aportad' in t.lower() or 'peso en oro' in t:
            print(f'  [{i:4d}] {t[:180]}')

print('\n=== References tail (last 30 non-empty paragraphs before Appendices) ===')
# Find references range
refs_start = None
apx_start = None
for i, p in enumerate(doc.paragraphs):
    if p.style and 'Heading 1' in p.style.name:
        if 'Referencias' in p.text:
            refs_start = i
        elif 'Apéndices' in p.text or 'Apendices' in p.text:
            apx_start = i
            break

if refs_start:
    print(f'  Referencias: paragraph {refs_start}, Apéndices: {apx_start}')
    # Print last ~10 refs
    for i in range(max(refs_start, apx_start - 15), apx_start):
        t = doc.paragraphs[i].text.strip()
        if t:
            print(f'  [{i:4d}] {t[:180]}')

print('\n=== Alternativa B rewrite check (461-473) ===')
for i, p in enumerate(doc.paragraphs):
    if 461 <= i <= 473:
        t = p.text.strip()[:150]
        if t:
            print(f'  [{i:4d}] {t}')
