#!/usr/bin/env python3
"""
Agrega resultados completos del Artículo 2 al informe v4 → genera v5.

Inserciones (siempre addnext sobre nodo fijo = orden correcto):
  1. Tabla real de 7 iteraciones después del párrafo 811 (Tabla 9.1)
  2. Figura confusion matrices después del párrafo 814
  3. Figura learning curve después del párrafo 817 (antes de h2 sig.)
  4. Figura flow correlation + SHAP después del párrafo 821
  5. Figura LOEO después del párrafo 821 (bloque de figuras apoyo)
  6. Figura calibration al final de la sección limitaciones
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

SRC  = Path('C:/Users/ayala/Downloads/Informe del proyecto_v4.docx')
OUT  = Path('C:/Users/ayala/Downloads/Informe del proyecto_v5.docx')
FIG  = Path('D:/pipeline_SVM/informe_proyecto/figures')

# Índices en el doc v4 (verificados)
IDX_TABLA     = 811   # párrafo "Tabla 9.1..."
IDX_AFTER_AUD = 814   # después del texto de audio-only → insertar confusion
IDX_AFTER_IMP = 817   # después del texto de meseta → insertar learning curve
IDX_AFTER_ROL = 821   # después del párrafo ESP32 → insertar flow+SHAP+LOEO
IDX_AFTER_LIM = 834   # después de limitaciones → insertar calibration

ITER_DATA = [
    (1, 2935,    0, 0.8644, 0.3343, 'Línea base Orejarena (reetiquetado [15/75])'),
    (2, 2935,    0, 0.8644, 0.3343, 'Línea base + coat (sin datos Art.2)'),
    (3, 3271,  336, 0.9228, 0.3538, '+ test39 (broca #4, 110 agujeros)'),
    (4, 4090, 1155, 0.8919, 0.3545, 'Reprocesado completo Art.2 (todos los tests)'),
    (5, 4228, 1293, 0.8936, 0.3540, '+ test53 (Dormer A114 #1, 44 agujeros)'),
    (6, 4381, 1446, 0.8936, 0.3513, '+ test50 (Dormer A100)'),
    (7, 4381, 1446, 0.8936, 0.3513, 'Audio-only → línea base multimodal'),
]
MULTI_ADJ   = 0.9537
MULTI_MACRO = 0.3535


# ── helpers ──────────────────────────────────────────────────────────────────

def insert_after(ref_node, new_elem):
    ref_node.addnext(new_elem)

def make_normal(doc, text, italic=False, size=None, color=None):
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def make_caption(doc, text):
    p = doc.add_paragraph(style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    return p

def make_figure(doc, img_path, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if Path(img_path).exists():
        run.add_picture(str(img_path), width=Inches(width))
    else:
        run.add_text(f'[imagen no encontrada: {img_path}]')
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p

def add_iterations_table(doc, ref_node):
    """Crea tabla Word con las 7 iteraciones y la fila multimodal."""
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'

    # Header
    hdr = tbl.rows[0].cells
    headers = ['Iter.', 'Muestras', 'Art.2', 'Adj. Acc. E3', 'Macro-F1', 'Descripción']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for it, n, n2, adj, mac, desc in ITER_DATA:
        row = tbl.add_row().cells
        vals = [str(it), str(n), str(n2) if n2 else '—',
                f'{adj:.4f}', f'{mac:.4f}', desc]
        for i, v in enumerate(vals):
            row[i].text = v
            for para in row[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 5 else WD_ALIGN_PARAGRAPH.LEFT

    # Multimodal final row (highlighted)
    row = tbl.add_row().cells
    vals = ['7*', '4 381', '1 446',
            f'{MULTI_ADJ:.4f}', f'{MULTI_MACRO:.4f}',
            'Multimodal completo (+ caudal + recubrimiento)']
    for i, v in enumerate(vals):
        row[i].text = v
        for para in row[i].paragraphs:
            run = para.runs[0] if para.runs else para.add_run(v)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x55, 0xAA)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 5 else WD_ALIGN_PARAGRAPH.LEFT

    ref_node.addnext(tbl._element)
    return tbl


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    doc = Document(str(SRC))
    paras = doc.paragraphs

    # ── 1. Tabla de iteraciones después del párrafo 811 ───────────────────────
    ref811 = paras[IDX_TABLA]
    # Insertar en orden inverso: caption primero (queda última), luego tabla
    cap_tab = make_caption(doc,
        'Tabla 9.1. Evolución de métricas del modelo SVM ordinal por iteración. '
        'Exactitud adyacente y Macro-F1 evaluados sobre la retención E3 (n=583). '
        'Fila 7* = variante multimodal completa (caudal + recubrimiento).')
    insert_after(ref811._element, cap_tab._element)
    add_iterations_table(doc, ref811._element)

    # ── 2. Figura confusion matrices después del párrafo 814 ──────────────────
    ref814 = paras[IDX_AFTER_AUD]
    cap_conf = make_caption(doc,
        'Figura 9.1. Matrices de confusión normalizadas por fila (retención E3, n=583). '
        '(A) Audio únicamente: adj=89,4 %. '
        '(B) Audio + recubrimiento: adj=88,5 %. '
        '(C) Multimodal completo: adj=95,4 % (+6,0 pp). '
        'La variante multimodal reduce los fallos desgastado→sin_desgaste de 62 a 27.')
    insert_after(ref814._element, cap_conf._element)
    fig_conf = make_figure(doc, FIG / 'fig_cap9_confusion.png', width=6.0)
    insert_after(ref814._element, fig_conf._element)

    # ── 3. Figura learning curve después del párrafo 817 ─────────────────────
    ref817 = paras[IDX_AFTER_IMP]
    cap_lc = make_caption(doc,
        'Figura 9.2. Curva de aprendizaje del modelo SVM ordinal. '
        'Eje X: número de muestras de entrenamiento acumuladas por iteración. '
        'Eje Y: exactitud adyacente en E3. La incorporación de los ensayos Art.2 '
        '(iter 3–6) estabiliza la métrica; el salto a la variante multimodal (iter 7*) '
        'rompe la meseta al añadir descriptores de caudal y recubrimiento.')
    insert_after(ref817._element, cap_lc._element)
    fig_lc = make_figure(doc, FIG / 'fig_cap9_lc.png', width=5.5)
    insert_after(ref817._element, fig_lc._element)

    # ── 4. Figuras de variables de apoyo después del párrafo 821 ─────────────
    ref821 = paras[IDX_AFTER_ROL]

    # LOEO (insertar último → quedará primero del bloque)
    cap_loeo = make_caption(doc,
        'Figura 9.5. Resultados de validación Leave-One-Experiment-Out (LOEO) '
        'sobre los siete experimentos de Orejarena. La variante multimodal mantiene '
        'exactitud adyacente ≥ 87 % en todos los experimentos excepto E1 y E2 '
        '(micrófonos dinámicos, mayor heterogeneidad de dominio).')
    insert_after(ref821._element, cap_loeo._element)
    fig_loeo = make_figure(doc, FIG / 'fig_cap9_loeo.png', width=5.5)
    insert_after(ref821._element, fig_loeo._element)

    # SHAP
    cap_shap = make_caption(doc,
        'Figura 9.4. Importancia de características (SHAP) para los clasificadores '
        'C₁ (¿hay desgaste?) y C₂ (¿es severo?). El caudal medio (flow_mean) aparece '
        'entre los cinco descriptores más relevantes de C₂, validando su papel como '
        'indicador de desgaste avanzado.')
    insert_after(ref821._element, cap_shap._element)
    fig_shap = make_figure(doc, FIG / 'fig_cap9_shap.png', width=5.5)
    insert_after(ref821._element, fig_shap._element)

    # Flow correlation
    cap_flow = make_caption(doc,
        'Figura 9.3. Correlación temporal entre caudal de refrigerante (YF-S201) '
        'y probabilidad de desgaste predicha (ventana 20 s) en el ensayo 53 '
        '(Dormer A114 #1, fractura en agujero 44). '
        'Correlación de Spearman ρ = −0,71 (p < 0,001): el caudal decae '
        'sistemáticamente a medida que el modelo detecta mayor desgaste.')
    insert_after(ref821._element, cap_flow._element)
    fig_flow = make_figure(doc, FIG / 'fig_cap9_flow.png', width=5.5)
    insert_after(ref821._element, fig_flow._element)

    # ── 5. Figura calibración después del párrafo 834 (limitaciones) ─────────
    ref834 = paras[IDX_AFTER_LIM]
    cap_calib = make_caption(doc,
        'Figura 9.6. Curvas de calibración de probabilidad por clase (retención E3). '
        'Los estados extremos (sin_desgaste, desgastado) están bien calibrados; '
        'la clase intermedia concentra mayor incertidumbre, coherente con la '
        'naturaleza borrosa del umbral de etiquetado [15/75].')
    insert_after(ref834._element, cap_calib._element)
    fig_calib = make_figure(doc, FIG / 'fig_cap9_calib.png', width=5.0)
    insert_after(ref834._element, fig_calib._element)

    doc.save(str(OUT))
    print(f'Guardado: {OUT}')
    print(f'Total párrafos v5: {len(doc.paragraphs)}')


if __name__ == '__main__':
    main()
