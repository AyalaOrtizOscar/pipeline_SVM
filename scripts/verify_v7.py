#!/usr/bin/env python3
"""Verify v7.docx structure after corrections."""
from docx import Document
from pathlib import Path
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = Path('C:/Users/ayala/Downloads/files/Informe_del_proyecto_v7.docx')
doc = Document(str(SRC))

print(f'Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}')

headings = []
for idx, p in enumerate(doc.paragraphs):
    if 'Heading' in (p.style.name if p.style else ''):
        headings.append((idx, p.style.name, p.text.strip()[:110]))

print(f'\nTotal headings: {len(headings)}')
print('\n--- All H1 and H2 headings ---')
for idx, style, txt in headings:
    if 'Heading 1' in style or 'Heading 2' in style:
        print(f'  [{idx:4d}] {style:15s} | {txt}')

# Check final paragraph
print(f'\n--- Last 10 paragraphs ---')
for p in doc.paragraphs[-10:]:
    txt = p.text.strip()[:100]
    style = p.style.name if p.style else ''
    print(f'  {style:20s} | {txt}')
